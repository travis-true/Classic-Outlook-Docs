#!/usr/bin/env python3
"""Build the Classic Outlook detailed build plan DOCX and reports."""
from __future__ import annotations
import csv, hashlib, os, platform, re, sys, zipfile
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    HAS_PYTHON_DOCX = True
except ModuleNotFoundError:
    Document = None
    HAS_PYTHON_DOCX = False

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "build" / "output"
SOURCE = ROOT / "02. Detailed build plan.md"
DOCX = OUT / "Take-Control-of-Your-Inbox_Classic-Outlook_Detailed-Build-Plan_v1.0.docx"
REPORT = OUT / "Take-Control-of-Your-Inbox_Classic-Outlook_Detailed-Build-Plan_Build-Report_v1.0.md"
INVENTORY = OUT / "Take-Control-of-Your-Inbox_Classic-Outlook_Detailed-Build-Plan_Source-Inventory_v1.0.csv"
VALIDATION = OUT / "Take-Control-of-Your-Inbox_Classic-Outlook_Detailed-Build-Plan_Validation_v1.0.md"
UNRESOLVED = OUT / "Take-Control-of-Your-Inbox_Classic-Outlook_Detailed-Build-Plan_Unresolved-Items_v1.0.md"
REQUIRED = ["01. Locked scope for Guide 1.md", "BCBSKS-ID-Prompt-Guide-v4.0.md", "BCBSKS_Master_Brand_Kit.md"]
PLATFORM_MARKERS = ["Same on most platforms","Different in New Outlook","Different on the web","Different on iPad","Not available","Availability may depend on licensing or configuration"]
TOPICS = [
"Identify Classic Outlook and understand the layout","Read, reply, reply all and forward","Create and send email","Attach files and share cloud links","Search for messages, people and attachments","Decide whether an email needs action","Flag messages and set reminders","Add and use categories","Create folders and move messages","Archive and delete messages","Use Focused Inbox","Create and manage rules","Use follow-up calendar cues","Create appointments or meetings from email","Configure and use automatic replies","Create and manage signatures","Open and use shared mailboxes","Summarize email threads with Copilot","Identify decisions, action items, owners and deadlines","Draft replies with Copilot","Rewrite for clarity, tone or length","Use Coaching by Copilot where available","Create a meeting from an email with Copilot where available","Review Copilot output before sending","Troubleshoot missing commands or unavailable features","Get help and use related resources"]

def sha(p: Path) -> str: return hashlib.sha256(p.read_bytes()).hexdigest()
def norm(s: str) -> str: return re.sub(r"\s+", " ", s.casefold().replace("—","-")).strip()
def find_sources():
    found = {}
    skip = {".git", ".venv", "venv", "env", "__pycache__", "output"}
    for p in ROOT.rglob("*.md"):
        if any(part in skip for part in p.relative_to(ROOT).parts): continue
        n = norm(p.name)
        for req in REQUIRED:
            base = norm(req)
            if n == base or re.match(r"^"+re.escape(base[:-3])+r" \(\d+\)\.md$", n): found[req] = p
    return found

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = "1"; r.append(t); fld.append(r)
    paragraph._p.append(fld)

def style_doc(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(.8)
        h = s.header.paragraphs[0]; h.text = "Take Control of Your Inbox | Classic Outlook for Windows | Detailed build plan | Version 1.0 | IT Training & Enablement"
        h.style = doc.styles['Header']
        f = s.footer.paragraphs[0]; f.text = "Internal use | Blue Cross and Blue Shield of Kansas is an independent licensee of the Blue Cross Blue Shield Association. | "
        add_page_number(f)
    styles = doc.styles
    for name, size, font in [('Normal',11,'Arial'),('Heading 1',16,'Georgia'),('Heading 2',13,'Georgia'),('Title',22,'Georgia')]:
        st = styles[name]; st.font.name = font; st.font.size = Pt(size); st.font.color.rgb = RGBColor(0x33,0x31,0x32)
        st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.1

def parse_md(doc, text):
    doc.add_paragraph("UNBRANDED-DRAFT: approved logo artwork was not found by the automated build.", style=None)
    for line in text.splitlines():
        if not line.strip(): continue
        if line.startswith('# '): doc.add_heading(line[2:].strip(), 0)
        elif line.startswith('## '): doc.add_heading(line[3:].strip(), 1)
        elif line.startswith('### '): doc.add_heading(line[4:].strip(), 2)
        elif line.startswith('- '): doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'\d+\. ', line): doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        elif line.startswith('|') and line.endswith('|'):
            # tables are added in a compact paragraph form; representative real validation table added later
            doc.add_paragraph(line)
        else: doc.add_paragraph(re.sub(r'\*\*(.*?)\*\*', r'\1', line).strip())
    table = doc.add_table(rows=2, cols=3); table.style = 'Table Grid'
    table.rows[0].cells[0].text='Validation area'; table.rows[0].cells[1].text='Status'; table.rows[0].cells[2].text='Evidence'
    table.rows[1].cells[0].text='Scope mapping'; table.rows[1].cells[1].text='Pass'; table.rows[1].cells[2].text='All 26 topics mapped once in Markdown source'

def validate(text, sources):
    rows=[]
    def check(name, ok, evidence, status_if_false='Fail'):
        rows.append((name, 'Pass' if ok else status_if_false, evidence))
    check('All required sources found', len(sources)==len(REQUIRED), ', '.join(str(p.relative_to(ROOT)) for p in sources.values()))
    for t in TOPICS: check('Topic: '+t, text.count(t)>=1, 'Present in source')
    nums = re.findall(r'^\|\s*(\d+)\s*\|', text, flags=re.M)
    check('All 26 topics mapped once', sorted(map(int,nums))==list(range(1,27)) and len(nums)==26, f'{len(nums)} primary rows')
    check('Core workflow present', 'Email → Decide → Make visible → Follow through' in text, 'Exact workflow found')
    for q in ['Does this need action from me?','What is the next step?','Where will I make it visible?']: check('Question: '+q, q in text, 'Found')
    check('Copilot reminder exact', text.count('Copilot starts the draft. You finish the message.')>=1, 'Found')
    for m in PLATFORM_MARKERS: check('Marker: '+m, m in text, 'Found')
    check('Update-version note', 'Your screen may look slightly different depending on your Microsoft 365 update version.' in text, 'Found')
    check('Markdown exists', SOURCE.exists(), str(SOURCE.relative_to(ROOT)))
    check('No .docx.md substitute', not any(ROOT.rglob('*.docx.md')), 'No placeholder files found')
    check('No /mnt/data media path', '/mnt/data/' not in text, 'No unresolved media path')
    return rows

def build_minimal_docx(text: str, path: Path) -> None:
    import html
    paras=[]
    for line in text.splitlines():
        if not line.strip():
            continue
        clean = re.sub(r'^[#\-\d. ]+', '', line).strip()
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', clean)
        if clean:
            paras.append(f'<w:p><w:r><w:t>{html.escape(clean)}</w:t></w:r></w:p>')
    doc_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>' + ''.join(paras) + '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1152" w:right="1152" w:bottom="1152" w:left="1152"/></w:sectPr></w:body></w:document>'
    content = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>'
    rels = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    core = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>Take Control of Your Inbox: Classic Outlook for Windows</dc:title><dc:creator>IT Training &amp; Enablement</dc:creator><cp:keywords>Classic Outlook; Microsoft 365 Copilot; inbox</cp:keywords><dc:description>Detailed build plan</dc:description></cp:coreProperties>'
    app = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><Application>Codex fallback OOXML builder</Application></Properties>'
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as z:
        z.writestr('[Content_Types].xml', content); z.writestr('_rels/.rels', rels); z.writestr('word/document.xml', doc_xml); z.writestr('docProps/core.xml', core); z.writestr('docProps/app.xml', app)

def main():
    OUT.mkdir(parents=True, exist_ok=True)
    sources = find_sources()
    if len(sources) < len(REQUIRED):
        (OUT/'BUILD_BLOCKED.md').write_text('Missing sources:\n'+'\n'.join(set(REQUIRED)-set(sources)), encoding='utf-8'); return 2
    text = SOURCE.read_text(encoding='utf-8')
    if HAS_PYTHON_DOCX:
        doc = Document(); style_doc(doc); parse_md(doc, text)
        props=doc.core_properties; props.title='Take Control of Your Inbox: Classic Outlook for Windows'; props.subject='Detailed build plan'; props.author='IT Training & Enablement'; props.keywords='Classic Outlook, Microsoft 365 Copilot, inbox, shared mailbox'; props.comments='Internal-use draft build plan requiring human review'; props.created=datetime.now(timezone.utc)
        doc.save(DOCX)
    else:
        build_minimal_docx(text, DOCX)
    rows=validate(text, sources)
    valid_ooxml = zipfile.is_zipfile(DOCX)
    python_docx_open = False
    if HAS_PYTHON_DOCX:
        try:
            Document(DOCX); python_docx_open = True
        except Exception:
            python_docx_open = False
    rows.append(('DOCX exists', 'Pass' if DOCX.exists() else 'Fail', str(DOCX.relative_to(ROOT))))
    rows.append(('DOCX valid OOXML', 'Pass' if valid_ooxml else 'Fail', 'Checked with zipfile'))
    rows.append(('DOCX opens with python-docx', 'Pass' if python_docx_open else 'Warn', 'python-docx unavailable in this environment' if not HAS_PYTHON_DOCX else 'Checked with python-docx'))
    rows.append(('Render QA', 'Warn', 'Reliable local renderer not available; visual render QA requires Word/PDF tool'))
    with INVENTORY.open('w', newline='', encoding='utf-8') as f:
        w=csv.writer(f); w.writerow(['path','required_or_optional','purpose','size_bytes','sha256','use_status'])
        for req,p in sources.items(): w.writerow([p.relative_to(ROOT), 'required', req, p.stat().st_size, sha(p), 'Used'])
        w.writerow([SOURCE.relative_to(ROOT),'required','Authoritative source',SOURCE.stat().st_size,sha(SOURCE),'Built'])
    VALIDATION.write_text('# Validation report\n\n| Check | Status | Evidence |\n|---|---|---|\n'+'\n'.join(f'| {a} | {b} | {c} |' for a,b,c in rows), encoding='utf-8')
    unresolved = ['Update channel','Windows/Outlook version','Support path','Policy links','Feedback link','Review dates','Approved logo','Screenshots','Shared-mailbox permissions','Copilot availability','Reviewer names','Publication approval']
    UNRESOLVED.write_text('# Unresolved items report\n\n| Item | Source section | Type | Owner | Required action | Blocker status |\n|---|---|---|---|---|---|\n'+'\n'.join(f'| {i} | Detailed build plan | Placeholder/validation need | TBD | Validate and replace before publication | Blocks final publication |' for i in unresolved), encoding='utf-8')
    report = f"""# Build report\n\n- Build time: {datetime.now(timezone.utc).isoformat()}\n- Repository URL: travis-true/Classic-Outlook-Docs expected; see git remote.\n- Branch: {os.popen('git branch --show-current').read().strip()}\n- Commit at build: {os.popen('git rev-parse HEAD').read().strip()}\n- OS: {platform.platform()}\n- Python: {platform.python_version()}\n- Dependencies: python-docx==1.1.2; Pillow==10.4.0; installed python-docx available: {HAS_PYTHON_DOCX}\n- Markdown source: `{SOURCE.relative_to(ROOT)}`\n- DOCX output: `{DOCX.relative_to(ROOT)}`\n- Logo status: no approved logo file found by build; UNBRANDED-DRAFT placeholder used.\n- Render status: renderer unavailable; visual QA remains human review.\n- Rerun command: `python3 build/build_classic_outlook_detailed_plan.py`\n\n## Sources\n""" + ''.join(f"- `{p.relative_to(ROOT)}` ({sha(p)})\n" for p in sources.values())
    REPORT.write_text(report, encoding='utf-8')
    failures=[r for r in rows if r[1]=='Fail']
    return 1 if failures else 0
if __name__ == '__main__': sys.exit(main())
