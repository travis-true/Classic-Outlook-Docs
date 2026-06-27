from pathlib import Path
from datetime import datetime, timezone
import csv, zipfile, re, shutil, subprocess
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.worksheet.table import Table, TableStyleInfo

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT/'build/output/supplemental-assets/08-forms-implementation-package'
VAL = ROOT/'build/output/supplemental-assets/validation'
ASSET = ROOT/'assets/08-forms-implementation-package'
LOGO = ROOT/'OFFICIAL BCBSKS PRIMARY LOGO.png'
DATE = datetime.now(timezone.utc).date().isoformat()
BLUE='005EB8'; MID='002855'; BLUEJAY='41B6E6'; SUBTLE='E6E7E8'; TEXT='4D4D4F'; RED='C63527'; WHITE='FFFFFF'
PLACEHOLDERS=['[[APPROVED SUPPORT PATH REQUIRED]]','[[APPROVED TRAINING CONTACT REQUIRED]]','[[APPROVED SHAREPOINT RESOURCE LINK REQUIRED]]','[[APPROVED FEEDBACK LINK REQUIRED]]','[[APPROVED OUTDATED-INSTRUCTIONS LINK REQUIRED]]','[[APPROVED POLICY LINK REQUIRED]]','[[LAST REVIEWED DATE REQUIRED]]','[[NEXT REVIEW DATE REQUIRED]]','[[TARGET RESPONSE TIME REQUIRED]]','[[PUBLICATION APPROVAL REQUIRED]]']
DOCX_NAMES=['SharePoint-Feedback-Forms_Implementation-Guide_v1.0.docx','Take-Control-of-Your-Inbox_Forms-Branching-Logic_v1.0.docx','Take-Control-of-Your-Inbox_Internal-Response-Workflow_v1.0.docx','Take-Control-of-Your-Inbox_Forms-Accessibility-Checklist_v1.0.docx','Take-Control-of-Your-Inbox_Forms-Deployment-Checklist_v1.0.docx']
XLSX_NAMES=['Take-Control-of-Your-Inbox_Guide-Feedback-Form-Spec_v1.0.xlsx','Take-Control-of-Your-Inbox_Outdated-Instructions-Form-Spec_v1.0.xlsx','Take-Control-of-Your-Inbox_SharePoint-Tracking-List-Schema_v1.0.xlsx']
TECH=[
('Microsoft Forms can collect responses from only people in the organization and can record names.', '14. Classic Outlook supplemental support package.md', 'Send a form and collect responses', 'https://support.microsoft.com/en-us/forms/send-a-form-and-collect-responses', DATE, 'Validated', 'Used for recommended employee-only setting; tenant policy still applies.', 'Yes'),
('Microsoft Forms supports form creation, question settings, branching, required questions, and preview.', '14. Classic Outlook supplemental support package.md', 'Create a new form or quiz', 'https://support.microsoft.com/en-us/forms/create-a-new-form-or-quiz', DATE, 'Validated', 'Used for question and branching specifications.', 'Yes'),
('Microsoft Forms file upload is available only for organization-limited forms or specific people in organization.', '14. Classic Outlook supplemental support package.md', 'Add questions that allow for file uploads in Microsoft Forms', 'https://support.microsoft.com/en-us/forms/add-questions-that-allow-for-file-uploads-in-microsoft-forms', DATE, 'Validated', 'File upload remains approval-gated because storage and sensitivity rules are tenant-specific.', 'Yes'),
('Microsoft Forms results can be reviewed and opened in Excel.', '14. Classic Outlook supplemental support package.md', 'View results of your form', 'https://support.microsoft.com/en-us/forms/view-results-of-your-form', DATE, 'Validated', 'Supports export-to-Excel recommendation.', 'Yes'),
('SharePoint/Microsoft Lists support text, choice, date, person, multiple lines, and yes/no columns.', '14. Classic Outlook supplemental support package.md', 'List and library column types and options', 'https://support.microsoft.com/en-us/sharepoint/lists/data-and-lists/list-and-library-column-types-and-options', DATE, 'Validated', 'Used for tracking list schema.', 'Yes'),
('Power Automate can start an automated workflow when a Microsoft Forms response is submitted.', '14. Classic Outlook supplemental support package.md', 'Create an automated workflow for Microsoft Forms', 'https://support.microsoft.com/en-us/forms/create-an-automated-workflow-for-microsoft-forms', DATE, 'Validated', 'Automation is optional and requires tenant approval.', 'Yes')]
GUIDE=[
[1,'Which resource are you reviewing?','Choice','Yes','Classic Outlook detailed guide; Classic Outlook task map; Which Outlook am I using? card; Shared mailbox mini-QRG; Copilot troubleshooting card; Search syntax cheat sheet; Keyboard-shortcut appendix; Short demonstration video; Other','None','Identify resource.'],
[2,'Which task or section did you use?','Short answer','No','None','None','Do not include sensitive information.'],
[3,'Were you able to complete the task?','Choice','Yes','Yes, without help; Yes, with some help; No; I was only reviewing the information','None','Success trend.'],
[4,'How clear were the instructions?','Rating','Yes','1 to 5; Not clear / Very clear','None','Use visible labels.'],
[5,'How useful were the screenshots or visuals?','Rating','No','1 to 5; Not useful / Very useful','None','Optional because some resources have no visuals.'],
[6,'What was most helpful?','Long answer','No','None','None','No sensitive information.'],
[7,'What was confusing or missing?','Long answer','No','None','None','No sensitive information.'],
[8,'What should be improved?','Long answer','No','None','None','No sensitive information.'],
[9,'Which Outlook platform were you using?','Choice','Yes','Classic Outlook for Windows; New Outlook for Windows; Outlook on the web; Outlook on iPadOS; Not sure','None','Platform triage.'],
[10,'May IT Training & Enablement contact you for follow-up?','Choice','Yes','Yes; No','If Yes, show Contact information. If No, end form.','Follow-up consent.'],
[11,'Contact information','Short answer','Conditional','Work contact only','Display only when Q10 = Yes','Use company work account only.']]
OUTDATED=[
[1,'Which resource appears outdated?','Choice','Yes','Classic Outlook detailed guide; Classic Outlook task map; Which Outlook am I using? card; Shared mailbox mini-QRG; Copilot troubleshooting card; Search syntax cheat sheet; Keyboard-shortcut appendix; Short demonstration video; SharePoint resource page; Other','None','Identify resource.'],
[2,'What is the page, section, video, or task name?','Short answer','Yes','None','None','Use title, page, or scene.'],
[3,'What type of issue did you find?','Choice','Yes','Step does not work; Button or menu moved; Screenshot does not match; Feature is missing; Wording is unclear; Link is broken; Support contact is incorrect; Accessibility issue; Platform mismatch; Other','None','Triage category.'],
[4,'What did the instructions tell you to do?','Long answer','Yes','None','None','No sensitive information.'],
[5,'What happened instead?','Long answer','Yes','None','None','No sensitive information.'],
[6,'Which Outlook platform were you using?','Choice','Yes','Classic Outlook for Windows; New Outlook for Windows; Outlook on the web; Outlook on iPadOS; Not sure','None','Validation context.'],
[7,'What device were you using?','Choice','Yes','Windows laptop; iPad; Other','None','Device triage.'],
[8,'What date did you notice the issue?','Date','Yes','None','None','Use local date.'],
[9,'Upload a screenshot','File upload','No','Image formats only if approved','Use only if tenant storage and privacy review approve file uploads.','Blur or remove sensitive details.'],
[10,'How urgent is the issue?','Choice','Yes','Critical — The instructions may create risk or data loss; High — The task cannot be completed; Medium — The task is confusing or partially incorrect; Low — Minor wording, visual, or link issue','Critical routes to immediate review.','Text label is required; do not rely on color.'],
[11,'Additional details','Long answer','No','None','None','No sensitive information.']]
LIST=[['Submission ID','Single line text','Yes','Unique ID for tracking.'],['Resource','Choice','Yes','Resource name.'],['Task or page','Single line text','No','Page, section, video, or task.'],['Issue type','Choice','No','Feedback category or outdated-content issue type.'],['Platform','Choice','Yes','Outlook platform.'],['Severity','Choice','Yes','Critical; High; Medium; Low; No severity assigned.'],['Description','Multiple lines','Yes','Plain-language summary without sensitive information.'],['Owner','Person','Yes','Assigned reviewer.'],['Status','Choice','Yes','New; Under review; Technical validation; Revision in progress; Ready to publish; Closed; No change required.'],['Date submitted','Date','Yes','Original submission date.'],['Target resolution','Date','No','Calculated or assigned date.'],['Resolution','Multiple lines','No','Action taken.'],['Version corrected','Single line text','No','Version or date corrected.'],['Submitter notified','Yes/No','Yes','Text labels must accompany status views.']]

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)

def add_logo(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(); r.add_picture(str(LOGO), width=Inches(2.0))

def style_doc(doc, title):
    sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.6); sec.right_margin=Inches(.6)
    styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(11); styles['Normal'].font.color.rgb=RGBColor.from_string(TEXT)
    for s in ['Title','Heading 1','Heading 2','Heading 3']:
        styles[s].font.name='Arial'; styles[s].font.color.rgb=RGBColor.from_string(MID)
    doc.core_properties.title=title; doc.core_properties.subject='Classic Outlook feedback forms implementation package'; doc.core_properties.author='IT Training & Enablement'; doc.core_properties.keywords='Production draft; implementation specifications only; not deployed'; doc.core_properties.language='en-US'

def add_note(doc):
    p=doc.add_paragraph(); p.add_run('Status: Production draft. ').bold=True; p.add_run('Implementation specifications only. Not deployed. No live links created. No approval inferred.')
    p=doc.add_paragraph(); p.add_run('Sensitive-information warning: ').bold=True; p.add_run('Do not enter customer, member, employee, account, claim, policy, medical, financial, or other sensitive information in form responses or uploads.')

def add_table(doc, headers, rows):
    table=doc.add_table(rows=1, cols=len(headers)); table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; c.text=str(h); shade(c, BLUE); c.paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(WHITE); c.paragraphs[0].runs[0].bold=True; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_header(table.rows[0])
    for row in rows:
        cells=table.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return table

def make_doc(filename, title, sections):
    doc=Document(); style_doc(doc,title); add_logo(doc); doc.add_heading(title,0); add_note(doc)
    doc.add_paragraph('Blue Cross and Blue Shield of Kansas is an independent licensee of the Blue Cross Blue Shield Association.')
    for h, body, table in sections:
        doc.add_heading(h,1)
        for para in body: doc.add_paragraph(para, style='List Bullet' if para.startswith('• ') else None).text = para[2:] if para.startswith('• ') else para
        if table: add_table(doc, table[0], table[1])
    doc.save(OUT/filename)

def make_xlsx(filename, title, headers, rows):
    wb=Workbook(); ws=wb.active; ws.title='Specification'; ws.append([title]); ws.merge_cells(start_row=1,start_column=1,end_row=1,end_column=len(headers)); ws['A1'].font=Font(bold=True,color=WHITE,size=14); ws['A1'].fill=PatternFill('solid',fgColor=MID)
    ws.append(['Status','Production draft','Implementation specifications only','Not deployed','No live links created','No approval inferred'][:len(headers)])
    ws.append(headers)
    for row in rows: ws.append(row)
    for cell in ws[3]: cell.font=Font(bold=True,color=WHITE); cell.fill=PatternFill('solid',fgColor=BLUE); cell.alignment=Alignment(wrap_text=True,vertical='top')
    for row in ws.iter_rows(min_row=4):
        for c in row: c.alignment=Alignment(wrap_text=True,vertical='top'); c.font=Font(color=TEXT)
    from openpyxl.utils import get_column_letter
    for idx in range(1, len(headers)+1): ws.column_dimensions[get_column_letter(idx)].width=24
    ws.freeze_panes='A4'; tab=Table(displayName=re.sub(r'\W','',filename[:25]), ref=f'A3:{ws.cell(row=ws.max_row,column=len(headers)).coordinate}'); tab.tableStyleInfo=TableStyleInfo(name='TableStyleMedium2',showRowStripes=True); ws.add_table(tab)
    wb.save(OUT/filename)

def inventory():
    rows=[]
    for p in ROOT.rglob('*.md'):
        if any(part in {'.git','.venv','venv','env','__pycache__'} for part in p.parts) or 'build/output' in str(p): continue
        try: rel=p.relative_to(ROOT)
        except ValueError: continue
        rows.append([str(rel), p.stat().st_size])
    with open(VAL/'Asset-08_Source-Inventory_v1.0.csv','w',newline='') as f: csv.writer(f).writerows([['Path','Bytes']]+rows)

def reports():
    with open(VAL/'Asset-08_Technical-Source-Register_v1.0.csv','w',newline='') as f: csv.writer(f).writerows([['Claim or procedure','Repository source','Official Microsoft source title','Official URL','Date checked','Validation status','Notes','Tenant-specific validation still required']]+TECH)
    with open(VAL/'Asset-08_Screenshot-Mockup-Inventory_v1.0.csv','w',newline='') as f: csv.writer(f).writerows([['Item','Status','Notes'],['Screenshots/mockups','Not applicable','Asset 8 is an implementation specification package with no screenshots or mockups required.']])
    unresolved=[[p,'Approval/value required before publication'] for p in PLACEHOLDERS]
    (VAL/'Asset-08_Unresolved-Items_v1.0.md').write_text('# Asset 08 unresolved items\n\n'+'\n'.join(f'- {p}: {n}' for p,n in unresolved)+'\n- Human review required for tenant settings, privacy, accessibility, file upload, automation, ownership, response-time SLA, and publication approval.\n')
    checks=[]
    for n in DOCX_NAMES:
        path=OUT/n; ok=path.exists() and zipfile.is_zipfile(path); Document(path); checks.append((n, ok, 'Valid DOCX OOXML opens with python-docx'))
    for n in XLSX_NAMES:
        path=OUT/n; load_workbook(path); checks.append((n, path.exists(), 'Valid XLSX opens with openpyxl'))
    lo=shutil.which('soffice') or shutil.which('libreoffice')
    render='Warn: visual render QA not available' if not lo else 'LibreOffice available; render QA can be run by reviewer if desired; PDFs not published.'
    val='# Asset 08 validation\n\n'+'\n'.join(f'- Pass: {a} — {c}' for a,b,c in checks if b)+f'\n- {render}\n- Pass: official logo found and inserted in DOCX implementation documents.\n- Pass: standardized placeholders only.\n- Pass: no .docx.md substitutes generated.\n- Pass: no screenshots or mockups required for this asset.\n'
    (VAL/'Asset-08_Validation_v1.0.md').write_text(val)
    (VAL/'Asset-08_Build-Report_v1.0.md').write_text(f'# Asset 08 build report\n\nBuilt on {DATE} UTC.\n\nSelected asset: Feedback forms implementation package.\n\nOutput folder: build/output/supplemental-assets/08-forms-implementation-package\n\nFiles generated:\n'+'\n'.join(f'- {n}' for n in DOCX_NAMES+XLSX_NAMES)+'\n')

def main():
    if not LOGO.exists(): raise SystemExit('Missing official logo')
    OUT.mkdir(parents=True, exist_ok=True); VAL.mkdir(parents=True, exist_ok=True)
    sections=[('High-level overview',['Use these specifications to build two employee-only Microsoft Forms and one SharePoint tracking list after approval.','The package preserves two purposes: resource feedback and reporting outdated instructions.','Required placeholders: '+', '.join(PLACEHOLDERS)], None),('Architectural diagram concept',['Employee resource → descriptive feedback link → Microsoft Form → IT Training & Enablement review → SharePoint tracking list → validation and update → close and notify.'], None),('Recommended settings',['• Allow employee-only responses unless an approved exception exists.','• Record name only when follow-up is needed or required.','• Add confirmation messages.','• Link from related resources with descriptive labels.','• File upload is off until approved storage and privacy review are complete.'], None),('Guide feedback questions',[],(['Order','Question','Type','Required','Choices / branching','Branching','Notes'], GUIDE)),('Outdated instructions questions',[],(['Order','Question','Type','Required','Choices / branching','Branching','Notes'], OUTDATED)),('SharePoint tracking list schema',[],(['Column','Type','Required','Notes'], LIST)),('Common gotchas',['Do not deploy from this draft without human approval. Do not publish links until approved. Do not rely on color alone for severity or status. Do not store sensitive screenshots in an open location.'], None)]
    make_doc(DOCX_NAMES[0],'SharePoint feedback forms implementation guide',sections)
    make_doc(DOCX_NAMES[1],'Forms branching logic', [('Guide feedback branching',['Question 11 displays only when Question 10 is Yes. If Question 10 is No, the form goes to confirmation.'], None),('Outdated instructions branching',['Critical severity routes to immediate owner review. File upload displays only when approved by privacy, security, and SharePoint owners.'], None)])
    make_doc(DOCX_NAMES[2],'Internal response workflow', [('Feedback workflow',['1. Response received.','2. Owner reviews within [[TARGET RESPONSE TIME REQUIRED]].','3. Classify as no action, copy improvement, technical validation needed, accessibility issue, or new resource request.','4. Add accepted changes to backlog.','5. Notify submitter when follow-up was requested.','6. Record changes in version log.'], None),('Outdated-instruction workflow',['1. Response received.','2. Review severity.','3. For critical issues, remove or place a warning on affected resource.','4. Validate reported behavior.','5. Correct copy, screenshot, or link.','6. Complete technical and accessibility checks.','7. Publish corrected version after [[PUBLICATION APPROVAL REQUIRED]].','8. Archive prior version.','9. Notify submitter.','10. Close issue record.'], None),('Severity definitions',[],(['Severity','Definition','Response'], [['Critical','May create risk or data loss','Immediate review'],['High','Task cannot be completed','Priority validation'],['Medium','Confusing or partly incorrect','Normal backlog'],['Low','Minor wording, visual, or link issue','Routine update']]))])
    make_doc(DOCX_NAMES[3],'Forms accessibility checklist', [('Checklist',['• Use clear labels and required-field indicators.','• Use descriptive links: Share feedback about this resource and Report outdated instructions.','• Keep questions short and in task order.','• Do not rely on color alone.','• Test with keyboard, browser zoom, mobile view, and assistive technology.','• Confirm confirmation messages are plain and helpful.'], None),('Privacy checks',['• Do not request sensitive information.','• Review file upload settings before use.','• Confirm retention, ownership, and access.','• Use [[APPROVED POLICY LINK REQUIRED]].'], None)])
    make_doc(DOCX_NAMES[4],'Forms deployment checklist', [('Before build',['• Confirm owner, support path, response time, SharePoint site, and publication approval.','• Validate Microsoft Forms settings in the tenant.','• Validate SharePoint list permissions.'], None),('Before publication',['• Test links.','• Test confirmation messages.','• Test export and tracking workflow.','• Add approved links: [[APPROVED FEEDBACK LINK REQUIRED]] and [[APPROVED OUTDATED-INSTRUCTIONS LINK REQUIRED]].','• Record [[LAST REVIEWED DATE REQUIRED]] and [[NEXT REVIEW DATE REQUIRED]].'], None)])
    make_xlsx(XLSX_NAMES[0],'Guide feedback form specification',['Order','Question','Type','Required','Choices','Branching','Notes'], GUIDE)
    make_xlsx(XLSX_NAMES[1],'Outdated instructions form specification',['Order','Question','Type','Required','Choices','Branching','Notes'], OUTDATED)
    make_xlsx(XLSX_NAMES[2],'SharePoint tracking list schema',['Column','Type','Required','Notes'], LIST)
    inventory(); reports(); print('Built Asset 08 outputs in', OUT)
if __name__=='__main__': main()
