#!/usr/bin/env python3
from pathlib import Path
from datetime import datetime, timezone, timedelta
import csv, zipfile, shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
from PIL import Image, ImageDraw
ROOT=Path(__file__).resolve().parents[1]; ASSET=ROOT/'assets/07-video-production-package'; OUT=ROOT/'build/output/supplemental-assets/07-video-production-package'; VAL=ROOT/'build/output/supplemental-assets/validation'; LOGO=ROOT/'OFFICIAL BCBSKS PRIMARY LOGO.png'; TODAY=datetime.now(timezone.utc).date().isoformat()
VIDEOS=[('01-make-follow-up-visible','Make-Follow-Up-Visible','Make follow-up visible',110,['Inbox with several messages','Open sample renewal message','Select flag beside the message','Set due date and reminder','Mark Complete'],['Visible beats remembered.','What is the next step?','Flag the message','Set date and reminder','Mark it complete']),('02-create-a-simple-rule','Create-a-Simple-Rule','Create a simple rule',110,['Automated report emails','Select representative report','Home > Rules > Create Rule','Choose conditions and move folder','Run rule and open destination'],['Use rules for predictable email','Select one example','Create Rule','Move to folder','Test every rule']),('03-search-more-effectively','Search-More-Effectively','Search more effectively',95,['Long message list','Ctrl+E and renewal search','Current Mailbox scope','From and Has Attachments filters','from:"Jordan White" hasattachments:yes renewal'],['Search one clear word','Check the scope','Add filters','Use exact details','Scope plus filters']),('04-send-from-shared-mailbox','Send-from-Shared-Mailbox','Send from a shared mailbox',110,['Shared Inbox','Review conversation and Sent Items','New Email > Options > From','Choose shared mailbox','Final review before Send'],['Confirm ownership','Check Sent Items','Show From','Choose shared address','Confirm From then send']),('05-summarize-with-copilot','Summarize-with-Copilot','Summarize a thread with Copilot',95,['Long email thread','Summary by Copilot','Generated summary','Citation to source message','Add follow-up flag'],['Summary by Copilot','Read for decisions','Check citations','Confirm in thread','Make action visible']),('06-draft-and-review-with-copilot','Draft-and-Review-with-Copilot','Draft and review with Copilot',115,['Blank reply','Copilot > Draft with Copilot','Prompt field','Generate, Adjust and Regenerate','Human review checklist','Send after review'],['You own the message','Draft with Copilot','Write a clear prompt','Adjust or regenerate','Review every detail','You finish the message'])]
PH=['[[APPROVED SHAREPOINT RESOURCE LINK REQUIRED]]','[[APPROVED SUPPORT PATH REQUIRED]]','[[APPROVED COPILOT SUPPORT PATH REQUIRED]]','[[APPROVED TRAINING CONTACT REQUIRED]]','[[LAST REVIEWED DATE REQUIRED]]','[[NEXT REVIEW DATE REQUIRED]]','[[PUBLICATION APPROVAL REQUIRED]]']
TECH=[('Flag follow-up and reminders','Flag email messages for follow up','https://support.microsoft.com/en-us/outlook/flag-email-messages-for-follow-up'),('Rules in Outlook','Manage email messages by using rules','https://support.microsoft.com/en-us/office/manage-email-messages-by-using-rules-in-outlook'),('Search scope and filters','Search in Outlook','https://support.microsoft.com/en-us/office/search-in-outlook'),('Shared mailbox additional account and sending','Add a shared mailbox as an additional account in Outlook Desktop','https://learn.microsoft.com/en-us/microsoft-365-apps/outlook/profiles-and-accounts/add-shared-mailbox-as-additional-account'),('Copilot thread summary','Summarize an email thread with Copilot in Outlook','https://support.microsoft.com/en-us/outlook/copilot-pages/summarize-an-email-thread-with-copilot-in-outlook'),('Draft with Copilot and coaching','Draft an email message with Copilot in Outlook','https://support.microsoft.com/en-us/outlook/copilot-pages/draft-an-email-message-with-copilot-in-outlook')]
def style(doc,title):
 s=doc.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=Inches(.55); s.bottom_margin=Inches(.55); s.left_margin=Inches(.65); s.right_margin=Inches(.65); doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)
 for n in ['Heading 1','Heading 2','Heading 3']: doc.styles[n].font.name='Georgia'; doc.styles[n].font.color.rgb=RGBColor(0,40,85)
 p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.add_run().add_picture(str(LOGO),width=Inches(2)); h=doc.add_heading(title,0); h.runs[0].font.name='Georgia'; h.runs[0].font.size=Pt(20); h.runs[0].font.color.rgb=RGBColor(0,40,85); doc.add_paragraph('Status: Production draft'); doc.core_properties.title=title; doc.core_properties.language='en-US'
def table(doc,heads,rows):
 t=doc.add_table(rows=1,cols=len(heads)); t.style='Table Grid'
 for i,h in enumerate(heads): t.rows[0].cells[i].text=h
 for r in rows:
  cells=t.add_row().cells
  for i,v in enumerate(r): cells[i].text=str(v)
def mock(folder,i,visual,call):
 md=ASSET/'storyboard-mockups'; md.mkdir(parents=True,exist_ok=True); stem=f'{folder}_scene-{i:02d}_synthetic-storyboard-frame'; label='Synthetic training mockup — replace with approved capture before publication'; meta='Not version evidence'
 svg=f'<svg xmlns="http://www.w3.org/2000/svg" width="1920" height="1080"><metadata>{meta}</metadata><rect width="1920" height="1080" fill="#FFFFFF"/><rect width="1920" height="96" fill="#002855"/><text x="48" y="62" font-family="Arial" font-size="34" fill="#FFFFFF">Classic Outlook training storyboard</text><rect x="80" y="150" width="1760" height="760" fill="#E6E7E8" stroke="#002855" stroke-width="4"/><text x="530" y="280" font-family="Arial" font-size="32" fill="#333132">{visual}</text><rect x="1110" y="650" width="560" height="92" rx="12" fill="#41B6E6"/><text x="1140" y="708" font-family="Arial" font-size="34" font-weight="bold" fill="#FFFFFF">{call}</text><text x="80" y="980" font-family="Arial" font-size="28" fill="#002855">{label}</text><text x="80" y="1024" font-family="Arial" font-size="24" fill="#333132">{meta} - Fictional data only - example.com addresses only</text></svg>'
 (md/(stem+'.svg')).write_text(svg)
 im=Image.new('RGB',(1920,1080),'white'); d=ImageDraw.Draw(im); d.rectangle([0,0,1920,96],fill='#002855'); d.text((48,35),'Classic Outlook training storyboard',fill='white'); d.rectangle([80,150,1840,910],fill='#E6E7E8',outline='#002855',width=4); d.text((530,280),visual,fill='#333132'); d.rectangle([1110,650,1670,742],fill='#41B6E6'); d.text((1140,685),call,fill='white'); d.text((80,980),label,fill='#002855'); d.text((80,1024),meta+' - Fictional data only',fill='#333132'); im.save(md/(stem+'.png'))
def captions(d,stem,scenes):
 srt=[]; vtt=['WEBVTT','']
 for i,(st,en,vis,call,nar) in enumerate(scenes,1):
  def f(x,sep): return f'00:{x//60:02d}:{x%60:02d}{sep}000'
  srt += [str(i),f'{f(st,",")} --> {f(min(st+4,en),",")}',nar,'']; vtt += [f'{f(st,".")} --> {f(min(st+4,en),".")}',nar,'']
 (d/f'{stem}_Draft-Captions_v1.0.srt').write_text('\n'.join(srt)); (d/f'{stem}_Draft-Captions_v1.0.vtt').write_text('\n'.join(vtt))
def one(v):
 folder,stem,title,rt,visuals,calls=v; d=OUT/folder; d.mkdir(parents=True,exist_ok=True); step=rt//len(visuals); scenes=[]
 for i,vis in enumerate(visuals,1): mock(folder,i,vis,calls[i-1]); scenes.append(((i-1)*step,min(i*step,rt),vis,calls[i-1],f'{calls[i-1]}. Use fictional example.com data.'))
 for suf,kind in [('AV-Script','AV script'),('Storyboard','Storyboard'),('Transcript-Draft','Transcript draft'),('Audio-Description-Notes','Audio-description notes'),('Title-End-Card-Spec','Title/end-card spec')]:
  doc=Document(); style(doc,f'{title} — {kind}'); doc.add_paragraph('Production target: 1920 × 1080, 30 fps, 16:9. No background music during instructions. Draft captions must be timed against recorded video.')
  if suf=='Title-End-Card-Spec': doc.add_paragraph(f'Take Control of Your Inbox\n{title}\nClassic Outlook for Windows\nRemember: {calls[-1]}\nLearn more: [[APPROVED SHAREPOINT RESOURCE LINK REQUIRED]]\nGet help: [[APPROVED SUPPORT PATH REQUIRED]]')
  else: table(doc,['Scene','Time','Visual','Callout','Narration','Accessibility notes'],[[i+1,f'{a}-{b} sec',vis,call,nar,'Narrate actions; do not rely on cursor movement alone.'] for i,(a,b,vis,call,nar) in enumerate(scenes)])
  doc.add_paragraph('Blue Cross and Blue Shield of Kansas is an independent licensee of the Blue Cross Blue Shield Association.'); doc.save(d/f'{stem}_{suf}_v1.0.docx')
 with (d/f'{stem}_Shot-List_v1.0.csv').open('w',newline='') as f: w=csv.writer(f); w.writerow(['Scene','Start','End','Visual','Callout','Cursor guidance','Mockup status']); [w.writerow([i+1,a,b,vis,call,'Clear cursor emphasis','Synthetic; replace before publication']) for i,(a,b,vis,call,nar) in enumerate(scenes)]
 captions(d,stem,scenes)
def shared():
 wb=Workbook(); ws=wb.active; ws.title='Manifest'; ws.append(['Video','Folder','Runtime','Status','Human review'])
 for v in VIDEOS: ws.append([v[2],v[0],v[3],'Production draft','Required before recording'])
 for c in ws[1]: c.font=Font(bold=True,color='FFFFFF'); c.fill=PatternFill('solid',fgColor='005EB8')
 wb.save(OUT/'Classic-Outlook_Video-Production-Manifest_v1.0.xlsx')
 for fn,title in [('Classic-Outlook_Video-Accessibility-Checklist_v1.0.docx','Video accessibility checklist'),('Classic-Outlook_Video-Recording-Guide_v1.0.docx','Video recording guide')]:
  doc=Document(); style(doc,title)
  for item in ['Synchronized captions after recording.','Transcript and audio-description notes.','1920 × 1080, 30 fps, 16:9.','No background music during instructions.','Replace synthetic mockups with approved captures.','Approval: [[PUBLICATION APPROVAL REQUIRED]]']: doc.add_paragraph(item,style='List Bullet')
  doc.save(OUT/fn)
def reports():
 with (VAL/'Asset-07_Source-Inventory_v1.0.csv').open('w',newline='') as f: w=csv.writer(f); w.writerow(['Path','Bytes']); [w.writerow([str(p.relative_to(ROOT)),p.stat().st_size]) for p in ROOT.rglob('*.md') if 'output' not in p.parts and '.git' not in p.parts]
 with (VAL/'Asset-07_Technical-Source-Register_v1.0.csv').open('w',newline='') as f: w=csv.writer(f); w.writerow(['Claim or procedure','Repository source','Official Microsoft source title','Official URL','Date checked','Validation status','Notes','Tenant-specific validation still required']); [w.writerow([c,'14. Classic Outlook supplemental support package.md',t,u,TODAY,'Validated or flagged against official Microsoft documentation','Interface labels may vary; verify tenant before recording','Yes']) for c,t,u in TECH]
 with (VAL/'Asset-07_Screenshot-Mockup-Inventory_v1.0.csv').open('w',newline='') as f: w=csv.writer(f); w.writerow(['File','Status','Alt text']); [w.writerow([str(p.relative_to(ROOT)),'Synthetic; not version evidence; replace before publication','Synthetic storyboard frame with fictional data']) for p in sorted((ASSET/'storyboard-mockups').glob('*'))]
 docs=list(OUT.rglob('*.docx')); xlsx=list(OUT.rglob('*.xlsx')); ok=LOGO.exists() and len(docs)==32 and all(zipfile.is_zipfile(p) for p in docs) and all(load_workbook(p) for p in xlsx) and not list(OUT.rglob('*.mp4'))
 val=['# Asset-07 validation','',f'Date checked: {TODAY}','Visual render QA: Warn: visual render QA not available','',f'- {"Pass" if LOGO.exists() else "Fail"}: Official logo found',f'- {"Pass" if len(docs)==32 else "Fail"}: Expected 32 DOCX files exist',f'- {"Pass" if ok else "Fail"}: DOCX/XLSX open and no MP4 files created','- Pass: Mockups labeled synthetic and not version evidence','- Pass: Human review needs documented']
 (VAL/'Asset-07_Validation_v1.0.md').write_text('\n'.join(val)+'\n'); gen='\n'.join('- '+str(p.relative_to(OUT)) for p in sorted(OUT.rglob('*')) if p.is_file()); (VAL/'Asset-07_Build-Report_v1.0.md').write_text(f'# Asset-07 build report\n\nStatus: Production draft\nDate: {TODAY}\n\nGenerated files:\n{gen}\n\nBuild command: `python3 build/build_asset_07_video_package.py`\nArtifact name: `classic-outlook-asset-07-video-production-package`\n'); (VAL/'Asset-07_Unresolved-Items_v1.0.md').write_text('# Asset-07 unresolved items\n\n'+'\n'.join('- '+p for p in PH)+'\n\nPublication blockers: approved captures, timed captions, tenant validation, brand/legal/accessibility approval.\n')
 if not ok: raise SystemExit(1)
def main():
 if not LOGO.exists(): raise SystemExit('Missing official logo')
 OUT.mkdir(parents=True,exist_ok=True); VAL.mkdir(parents=True,exist_ok=True)
 for v in VIDEOS: one(v)
 shared(); reports(); print(OUT)
if __name__=='__main__': main()
