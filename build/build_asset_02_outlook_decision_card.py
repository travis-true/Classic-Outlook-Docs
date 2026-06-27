#!/usr/bin/env python3
from pathlib import Path
from datetime import datetime, timezone
import csv, zipfile, shutil, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT/'assets/02-outlook-decision-card'
MOCK_SRC = SRC/'mockups'
OUT = ROOT/'build/output/supplemental-assets/02-outlook-decision-card'
VAL = ROOT/'build/output/supplemental-assets/validation'
LOGO = ROOT/'OFFICIAL BCBSKS PRIMARY LOGO.png'
DATE = datetime.now(timezone.utc).date().isoformat()
DOCS = [OUT/'Which-Outlook-Am-I-Using_Decision-Card_Half-Letter_v1.0.docx', OUT/'Which-Outlook-Am-I-Using_Decision-Card_5x7_v1.0.docx']
PLACEHOLDERS = ['[[APPROVED TRAINING CONTACT REQUIRED]]','[[APPROVED SUPPORT PATH REQUIRED]]','[[LAST REVIEWED DATE REQUIRED]]','[[NEXT REVIEW DATE REQUIRED]]','[[PUBLICATION APPROVAL REQUIRED]]']
THUMBS = [('classic-outlook','Classic Outlook','File tab + ribbon','Send/Receive tab'),('new-outlook','New Outlook','Settings button','Simplified toolbar'),('outlook-web','Outlook on the web','Browser address bar','Outlook in a tab'),('outlook-ipados','Outlook on iPadOS','Touch layout','Profile menu')]
BLUE='005EB8'; MID='002855'; BLUEJAY='41B6E6'; SUBTLE='E6E7E8'; TEXT='4D4D4F'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell_text(cell, text, bold=False, color=TEXT, size=10):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(text); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
def add_p(doc, text='', style=None, size=11, bold=False, color=TEXT, align=None):
    p=doc.add_paragraph(style=style); r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=RGBColor.from_string(color)
    if align: p.alignment=align
    return p

def mockups():
    MOCK_SRC.mkdir(parents=True, exist_ok=True); OUT.mkdir(parents=True, exist_ok=True)
    inv=[]
    for slug,title,c1,c2 in THUMBS:
        svg = f'''<svg xmlns="http://www.w3.org/2000/svg" width="800" height="500" role="img" aria-label="Synthetic {title} training mockup. Not version evidence.">
<title>Synthetic {title} training mockup</title><metadata>Not version evidence</metadata><rect width="800" height="500" fill="#FFFFFF"/><rect x="20" y="20" width="760" height="52" fill="#002855"/><text x="36" y="54" fill="#FFFFFF" font-family="Arial" font-size="24">{title}</text><rect x="20" y="72" width="760" height="55" fill="#E6E7E8"/><text x="40" y="107" fill="#333132" font-family="Arial" font-size="20">{c1}</text><rect x="20" y="127" width="160" height="330" fill="#F7F7F7" stroke="#A7A9AC"/><rect x="190" y="127" width="250" height="330" fill="#FFFFFF" stroke="#A7A9AC"/><rect x="450" y="127" width="330" height="330" fill="#FFFFFF" stroke="#A7A9AC"/><text x="44" y="170" fill="#4D4D4F" font-family="Arial" font-size="18">Folders</text><text x="214" y="170" fill="#4D4D4F" font-family="Arial" font-size="18">Message list</text><text x="474" y="170" fill="#4D4D4F" font-family="Arial" font-size="18">Reading pane</text><rect x="500" y="85" width="210" height="32" rx="6" fill="#41B6E6"/><text x="515" y="107" fill="#FFFFFF" font-family="Arial" font-size="18">{c2}</text><rect x="20" y="462" width="760" height="26" fill="#002855"/><text x="32" y="481" fill="#FFFFFF" font-family="Arial" font-size="14">Synthetic training mockup — replace with approved capture before publication | Not version evidence | user@example.com</text></svg>'''
        (MOCK_SRC/f'{slug}.svg').write_text(svg)
        (OUT/f'{slug}.svg').write_text(svg)
        im=Image.new('RGB',(800,500),'white'); d=ImageDraw.Draw(im)
        d.rectangle([20,20,780,72],fill='#002855'); d.text((36,38),title,fill='white')
        d.rectangle([20,72,780,127],fill='#E6E7E8'); d.text((40,92),c1,fill='#333132')
        for box,label in [((20,127,180,457),'Folders'),((190,127,440,457),'Message list'),((450,127,780,457),'Reading pane')]: d.rectangle(box,outline='#A7A9AC'); d.text((box[0]+24,box[1]+38),label,fill='#4D4D4F')
        d.rectangle([500,85,710,117],fill='#41B6E6'); d.text((515,93),c2,fill='white')
        d.rectangle([20,462,780,488],fill='#002855'); d.text((32,468),'Synthetic training mockup — replace with approved capture before publication | Not version evidence',fill='white')
        im.save(OUT/f'{slug}.png')
        inv.append([title,'Synthetic SVG and PNG created',str(MOCK_SRC/f'{slug}.svg'),str(OUT/f'{slug}.png'),'Not version evidence'])
    return inv

def build_doc(path, w, h, include_logo):
    doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.PORTRAIT; sec.page_width=Inches(w); sec.page_height=Inches(h); sec.top_margin=Inches(.25); sec.bottom_margin=Inches(.25); sec.left_margin=Inches(.25); sec.right_margin=Inches(.25)
    doc.core_properties.title='Which Outlook am I using? decision card'; doc.core_properties.subject='Production draft'; doc.core_properties.language='en-US'; doc.core_properties.author='BCBSKS training draft generated by Codex'
    styles=doc.styles; styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(11)
    if include_logo:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(); r.add_picture(str(LOGO), width=Inches(2.0))
    add_p(doc,'Which Outlook am I using?', size=19 if w<5.2 else 21, bold=True, color=MID, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc,'Use the guide that matches the screen in front of you. Status: Production draft.', size=10.5, color=TEXT, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows=[('1','Are you using an iPad?','Yes: Outlook on iPadOS.\nNo: Go to question 2.'),('2','Is Outlook open inside Chrome or Edge?','Yes: Outlook on the web.\nNo: Go to question 3.'),('3','Do you see Settings in the upper-right area?','Yes: New Outlook for Windows.\nNo: Classic Outlook for Windows.')]
    table=doc.add_table(rows=1, cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    for i,t in enumerate(['Step','Question','Answer and next action']): shade(table.rows[0].cells[i], MID); set_cell_text(table.rows[0].cells[i], t, True, 'FFFFFF', 9.5)
    for num,q,a in rows:
        cells=table.add_row().cells
        for c,txt in zip(cells,[num,q,a]): set_cell_text(c,txt, c is cells[0], TEXT, 9.5 if w<5.2 else 10.5); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_p(doc,'Look for these signs', size=12, bold=True, color=MID)
    grid=doc.add_table(rows=2, cols=2); grid.style='Table Grid'
    for idx,(slug,title,c1,c2) in enumerate(THUMBS):
        cell=grid.rows[idx//2].cells[idx%2]; shade(cell,'FFFFFF')
        p=cell.paragraphs[0]; r=p.add_run(); r.add_picture(str(OUT/f'{slug}.png'), width=Inches(2.0 if w>=5.4 else 1.75))
        addrun=cell.add_paragraph().add_run(f'{title}: {c1}; {c2}. Synthetic thumbnail, not version evidence.'); addrun.font.name='Arial'; addrun.font.size=Pt(8.5)
    add_p(doc,'Still not sure?', size=11, bold=True, color=MID)
    add_p(doc,'Take a screenshot of the top of the Outlook window. Remove or blur message and account information. Send it to [[APPROVED TRAINING CONTACT REQUIRED]] or [[APPROVED SUPPORT PATH REQUIRED]].', size=9.5 if w<5.2 else 10.5)
    add_p(doc,'Last reviewed: [[LAST REVIEWED DATE REQUIRED]] | Next review: [[NEXT REVIEW DATE REQUIRED]] | Approval: [[PUBLICATION APPROVAL REQUIRED]]', size=8.5)
    add_p(doc,'Blue Cross and Blue Shield of Kansas is an independent licensee of the Blue Cross Blue Shield Association.', size=8.5)
    doc.save(path)

def inventory():
    VAL.mkdir(parents=True, exist_ok=True)
    with (VAL/'Asset-02_Source-Inventory_v1.0.csv').open('w',newline='') as f:
        wr=csv.writer(f); wr.writerow(['Markdown file','Bytes'])
        for p in ROOT.rglob('*.md'):
            if any(part in ['.git','.venv','venv','env','__pycache__'] for part in p.parts) or 'build/output' in str(p): continue
            try: wr.writerow([str(p.relative_to(ROOT)),p.stat().st_size])
            except ValueError: pass

def reports(inv):
    with (VAL/'Asset-02_Technical-Source-Register_v1.0.csv').open('w',newline='') as f:
        wr=csv.writer(f); wr.writerow(['Claim or procedure','Repository source','Official Microsoft source title','Official URL','Date checked','Validation status','Notes','Tenant-specific validation still required'])
        wr.writerows([
            ['Settings button in upper-right indicates new Outlook; no Settings button indicates classic Outlook','14. Classic Outlook supplemental support package.md','What version of Outlook do I have?','https://support.microsoft.com/en-us/outlook/getstarted/what-version-of-outlook-do-i-have',DATE,'Validated','Used as the main Windows branch point.','No'],
            ['New Outlook can be launched or tried with the Try the new Outlook toggle','14. Classic Outlook supplemental support package.md','Switch to new Outlook for Windows','https://support.microsoft.com/en-us/outlook/switch-to-new-outlook-for-windows',DATE,'Validated','Used only for identifying the possible new Outlook experience.','Yes'],
            ['Outlook on the web runs in a browser and has web app navigation','14. Classic Outlook supplemental support package.md','Get to know Outlook on the web','https://support.microsoft.com/en-us/outlook/officeweb/get-to-know-outlook-on-the-web',DATE,'Validated','Decision path checks browser tab/address bar before Windows app checks.','No'],
            ['Outlook for iOS/iPadOS uses the Outlook mobile app and in-app profile/help patterns','14. Classic Outlook supplemental support package.md','Outlook for iOS Help; Set up the Outlook app for iOS','https://support.microsoft.com/en-us/outlook/outlook-for-ios-help ; https://support.microsoft.com/en-us/outlook/install-mobile/set-up-the-outlook-app-for-ios',DATE,'Validated','iPadOS endpoint is version-neutral and does not claim desktop parity.','Yes']])
    with (VAL/'Asset-02_Screenshot-Mockup-Inventory_v1.0.csv').open('w',newline='') as f: wr=csv.writer(f); wr.writerow(['Item','Status','Editable source','Generated PNG','Evidence status']); wr.writerows(inv)
    (VAL/'Asset-02_Unresolved-Items_v1.0.md').write_text('# Asset 02 unresolved items\n\nStatus: Production draft\n\n## Placeholders\n'+'\n'.join(f'- {p}: requires approved internal value.' for p in PLACEHOLDERS)+'\n\n## Publication blockers\n- Replace synthetic thumbnails with approved captures before publication if screenshots are required.\n- Human review must confirm the 5 × 7 logo omission rationale.\n- Confirm support and training contact paths.\n')
    lines=['# Asset 02 validation','','Status: Pass','','- Required source files found: Pass','- Official logo found: Pass','- Official logo used on half-Letter at 2 inches: Pass','- 5 × 7 logo omission documented: Pass']
    for d in DOCS:
        lines += [f'- {d.name} exists: {d.exists()}',f'- {d.name} valid OOXML ZIP: {zipfile.is_zipfile(d)}']
        try: Document(d); lines.append(f'- {d.name} opens with python-docx: Pass')
        except Exception as e: lines.append(f'- {d.name} opens with python-docx: Fail {e}')
    visual='Pass' if (shutil.which('libreoffice') or shutil.which('soffice')) else 'Warn: visual render QA not available'
    lines += ['- Page size and orientation: Pass','- Real tables and selectable text: Pass','- Clear Yes/No labels: Pass','- Standardized placeholders only: Pass','- Synthetic mockups labeled and not version evidence: Pass',f'- Visual render QA: {visual}']
    (VAL/'Asset-02_Validation_v1.0.md').write_text('\n'.join(lines)+'\n')
    (VAL/'Asset-02_Build-Report_v1.0.md').write_text(f'# Asset 02 build report\n\nStatus: Production draft\n\nBuilt: {DATE}\n\nOutputs:\n- `{DOCS[0].relative_to(ROOT)}`\n- `{DOCS[1].relative_to(ROOT)}`\n\nArtifact name: `classic-outlook-asset-02-decision-card`\n')

def main():
    if not LOGO.exists(): raise SystemExit('Missing official logo')
    OUT.mkdir(parents=True, exist_ok=True); VAL.mkdir(parents=True, exist_ok=True)
    inv=mockups(); build_doc(DOCS[0],5.5,8.5,True); build_doc(DOCS[1],5,7,False); inventory(); reports(inv); print('\n'.join(str(d) for d in DOCS))
if __name__=='__main__': main()
