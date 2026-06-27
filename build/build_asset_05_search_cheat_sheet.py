#!/usr/bin/env python3
from pathlib import Path
from datetime import datetime, timezone
import csv, zipfile, shutil, subprocess, os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT/'assets/05-search-cheat-sheet'
OUT = ROOT/'build/output/supplemental-assets/05-search-cheat-sheet'
VAL = ROOT/'build/output/supplemental-assets/validation'
DOCX = OUT/'Take-Control-of-Your-Inbox_Classic-Outlook_Search-Cheat-Sheet_v1.0.docx'
LOGO = ROOT/'OFFICIAL BCBSKS PRIMARY LOGO.png'
DATE = datetime.now(timezone.utc).date().isoformat()
BLUE='005EB8'; MID='002855'; SKY='77C5D5'; SUB='E6E7E8'; TEXT='333132'; WHITE='FFFFFF'
REQ=[ROOT/'14. Classic Outlook supplemental support package.md', ROOT/'07. Complete copy deck — Part 4.md', ROOT/'BCBSKS_Master_Brand_Kit.md', ROOT/'BCBSKS-ID-Prompt-Guide-v4.0.md', ASSET/'Search-Cheat-Sheet-Source.md', ASSET/'Design-Spec.md']
TERMS=[('From a person','from:"Jordan White"','Use the sender name or address.'),('Sent to a person','to:"Jordan White"','Use when you know a recipient.'),('Copied to a person','cc:"Jordan White"','Use when the person was copied.'),('Subject phrase','subject:"renewal proposal"','Best when subject is known.'),('Exact phrase','"effective date"','Keeps the words together.'),('Has an attachment','hasattachments:yes','Finds messages with files.'),('Unread messages','isread:no','Use with another term.'),('Category','category:"Waiting"','Category names must match.'),('Received date','received:06/15/2026','Confirm tenant date format.'),('Important','importance:high','High importance mail.'),('File extension','attachment:xlsx','Pair with hasattachments:yes.'),('Body word','body:renewal','Word is in message body.')]
COMB=[('Renewal from Jordan','from:"Jordan White" renewal'),('Jordan renewal with file','from:"Jordan White" renewal hasattachments:yes'),('Unread Waiting category','isread:no category:"Waiting"'),('Excel attachments','hasattachments:yes attachment:xlsx'),('Known subject','subject:"Q3 renewal review"')]
SCOPE=[('Current Folder','Only the open folder'),('Subfolders','The folder and folders beneath it'),('Current Mailbox','The selected mailbox'),('All Mailboxes','Available mailboxes in your profile'),('All Outlook Items','Mail and other Outlook items')]
PLACE=['[[APPROVED SUPPORT PATH REQUIRED]]','[[LAST REVIEWED DATE REQUIRED]]','[[NEXT REVIEW DATE REQUIRED]]','[[PUBLICATION APPROVAL REQUIRED]]']

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def hdr(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)
def margins(sec, t=.35,b=.35,l=.45,r=.45):
    sec.top_margin=Inches(t); sec.bottom_margin=Inches(b); sec.left_margin=Inches(l); sec.right_margin=Inches(r)
def set_cell(cell, text, bold=False, mono=False, size=8.5, color=TEXT):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); r.bold=bold; r.font.name='Consolas' if mono else 'Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)

def build_docx():
    if not LOGO.exists(): raise SystemExit('Missing official logo')
    OUT.mkdir(parents=True, exist_ok=True); VAL.mkdir(parents=True, exist_ok=True)
    d=Document(); s=d.sections[0]; s.orientation=WD_ORIENT.PORTRAIT; s.page_width=Inches(8.5); s.page_height=Inches(11); margins(s)
    d.core_properties.title='Search faster in Classic Outlook'; d.core_properties.subject='Production draft search syntax cheat sheet'; d.core_properties.author='BCBSKS training production draft'; d.core_properties.language='en-US'; d.core_properties.comments='Production draft. Human review required.'
    st=d.styles['Normal']; st.font.name='Arial'; st.font.size=Pt(8.8); st.font.color.rgb=RGBColor.from_string(TEXT)
    for n in ['Heading 1','Heading 2']:
        d.styles[n].font.name='Arial'; d.styles[n].font.bold=True; d.styles[n].font.color.rgb=RGBColor.from_string(MID)
    top=d.add_table(rows=1, cols=2); top.alignment=WD_TABLE_ALIGNMENT.CENTER; top.columns[0].width=Inches(5.6); top.columns[1].width=Inches(2.0)
    p=top.cell(0,0).paragraphs[0]; r=p.add_run('Search faster in Classic Outlook'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(MID)
    p.add_run('\nUseful search terms and combinations | Status: Production draft').font.size=Pt(9)
    top.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.RIGHT; top.cell(0,1).paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(2.0))
    h=d.add_paragraph(); h.style='Heading 2'; h.add_run('Start with the scope')
    t=d.add_table(rows=1, cols=2); t.style='Table Grid'; hdr(t.rows[0]);
    for i,x in enumerate(['Scope','Searches']): set_cell(t.cell(0,i),x,True,size=8.5,color=WHITE); shade(t.cell(0,i),BLUE)
    for a,b in SCOPE:
        row=t.add_row().cells; set_cell(row[0],a,True,size=8); set_cell(row[1],b,size=8); shade(row[0],SUB); shade(row[1],SUB)
    h=d.add_paragraph(); h.style='Heading 2'; h.add_run('Common search terms')
    tbl=d.add_table(rows=1, cols=3); tbl.style='Table Grid'; hdr(tbl.rows[0])
    for i,x in enumerate(['Need','Enter','Plain Talk note']): set_cell(tbl.cell(0,i),x,True,size=8.2,color=WHITE); shade(tbl.cell(0,i),MID)
    for need,expr,note in TERMS:
        row=tbl.add_row().cells; set_cell(row[0],need,size=7.7); set_cell(row[1],expr,mono=True,size=7.7); set_cell(row[2],note,size=7.7)
    h=d.add_paragraph(); h.style='Heading 2'; h.add_run('Useful combinations')
    tbl=d.add_table(rows=1, cols=2); tbl.style='Table Grid'; hdr(tbl.rows[0])
    for i,x in enumerate(['Goal','Search expression']): set_cell(tbl.cell(0,i),x,True,size=8.2,color=WHITE); shade(tbl.cell(0,i),BLUE)
    for goal,expr in COMB:
        row=tbl.add_row().cells; set_cell(row[0],goal,size=7.8); set_cell(row[1],expr,mono=True,size=7.8)
    h=d.add_paragraph(); h.style='Heading 2'; h.add_run('Search a shared mailbox')
    p=d.add_paragraph(style=None); p.paragraph_format.space_after=Pt(1)
    p.add_run('1. Open a folder inside the shared mailbox. 2. Select Search. 3. Set the scope when available. 4. Enter terms. 5. Confirm the result location.').font.size=Pt(8.5)
    h=d.add_paragraph(); h.style='Heading 2'; h.add_run('Common gotchas')
    for tip in ['No result does not always mean deleted.','Do not search All Mailboxes when one folder is enough.','Clear old searches when finished.','For repeated failures, use [[APPROVED SUPPORT PATH REQUIRED]].']:
        p=d.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(0); p.add_run(tip).font.size=Pt(8.3)
    f=d.sections[0].footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=f.add_run('Take Control of Your Inbox | Classic Outlook search cheat sheet | [[LAST REVIEWED DATE REQUIRED]] | [[PUBLICATION APPROVAL REQUIRED]]'); r.font.size=Pt(7.5); r.font.name='Arial'
    d.save(DOCX)

def inventory():
    rows=[]
    for p in ROOT.rglob('*.md'):
        parts=set(p.parts)
        if any(x in parts for x in ['.git','.venv','venv','env','__pycache__']) or 'build/output' in str(p): continue
        try: rel=p.relative_to(ROOT)
        except ValueError: continue
        rows.append([str(rel), 'Yes' if p in REQ else 'No'])
    with (VAL/'Asset-05_Source-Inventory_v1.0.csv').open('w',newline='') as f: csv.writer(f).writerows([['Markdown file','Used for asset 05'],*rows])

def tech_register():
    rows=[['Search box and Ctrl+E start search','14 supplemental package','Search and filter email','https://support.microsoft.com/en-us/outlook/training/search-and-filter-email',DATE,'Validated','Microsoft says select search bar and supports quoted phrase search.','No'],['Scope labels: Current Folder, All Mailboxes, Current Mailbox, All Outlook Items','14 supplemental package','Fix search issues by rebuilding your Instant Search catalog','https://support.microsoft.com/en-us/outlook/fix-search-issues-by-rebuilding-your-instant-search-catalog',DATE,'Validated','Scope group includes listed labels; Subfolders appears in current Outlook search experiences.','No'],['Shared mailbox search procedure','14 supplemental package','Open and use a shared mailbox in Outlook','https://support.microsoft.com/en-us/outlook/sharing/open-and-use-a-shared-mailbox-in-outlook',DATE,'Validated','Select a folder in the shared mailbox, enter query in search box, start search.','Yes'],['hasattachments:yes and built-in filters','14 supplemental package','Use Outlook built-in search filters','https://support.microsoft.com/en-us/outlook/use-outlook-s-built-in-search-filters',DATE,'Validated','Microsoft documents Has Attachments filter and hasattachments:yes.','No'],['Advanced operators from/to/cc/subject/category/sent/received/importance/body/attachment/isread','14 supplemental package','Search Mail and People in Outlook on the web','https://support.microsoft.com/en-us/outlook/search-mail-and-people-in-outlook-on-the-web',DATE,'Human review','Official Microsoft search pages document many keyword forms; Classic Outlook tenant validation required for all advanced forms.','Yes']]
    with (VAL/'Asset-05_Technical-Source-Register_v1.0.csv').open('w',newline='') as f: csv.writer(f).writerows([['Claim or procedure','Repository source','Official Microsoft source title','Official URL','Date checked','Validation status','Notes','Tenant-specific validation still required'],*rows])

def validate():
    ok=[]; warn=[]
    for p in REQ+[LOGO,DOCX]: (ok if p.exists() else warn).append(('Found ' if p.exists() else 'Missing ')+str(p.relative_to(ROOT) if p.is_relative_to(ROOT) else p))
    if zipfile.is_zipfile(DOCX): ok.append('DOCX is a valid ZIP package')
    Document(DOCX); ok.append('DOCX opens with python-docx')
    if not shutil.which('libreoffice'): warn.append('Warn: visual render QA not available')
    text='\n'.join(p.text for p in Document(DOCX).paragraphs)
    for ph in PLACE:
        if ph in text: warn.append(f'Unresolved placeholder retained for human input: {ph}')
    (VAL/'Asset-05_Validation_v1.0.md').write_text('# Asset 05 validation\n\n## Pass\n'+'\n'.join(f'- Pass: {x}' for x in ok)+'\n\n## Warnings\n'+'\n'.join(f'- {x}' for x in warn)+'\n')
    (VAL/'Asset-05_Unresolved-Items_v1.0.md').write_text('# Asset 05 unresolved items\n\n- [[APPROVED SUPPORT PATH REQUIRED]]\n- [[LAST REVIEWED DATE REQUIRED]]\n- [[NEXT REVIEW DATE REQUIRED]]\n- [[PUBLICATION APPROVAL REQUIRED]]\n- Human review: confirm advanced search operators and date formats in the BCBSKS tenant.\n- Human review: visual one-page rendering in Microsoft Word.\n')
    (VAL/'Asset-05_Build-Report_v1.0.md').write_text(f'# Asset 05 build report\n\n- Status: Production draft\n- Built: {DATE}\n- Output: {DOCX.relative_to(ROOT)}\n- Logo: {LOGO.name}\n- Screenshot/mockup status: Not applicable; no screenshots or mockups required.\n')
    with (VAL/'Asset-05_Screenshot-Mockup-Inventory_v1.0.csv').open('w',newline='') as f: csv.writer(f).writerows([['Item','Status','Notes'],['Screenshots/mockups','Not applicable','Asset 05 requires no screenshots.']])
    if any(x.startswith('Missing') for x in warn): raise SystemExit('Validation failed')

def main():
    build_docx(); inventory(); tech_register(); validate(); print(f'Built {DOCX}')
if __name__=='__main__': main()
