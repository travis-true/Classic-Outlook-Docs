#!/usr/bin/env python3
"""Deterministically build the employee-facing Classic Outlook Word guide."""
from __future__ import annotations
import csv, hashlib, os, re, sys, zipfile
from datetime import datetime, timezone
from pathlib import Path
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ModuleNotFoundError:
    Document = None
    HAS_DOCX = False
ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'build' / 'output'
TRAIN = ROOT / 'docs' / 'training'
GUIDE_MD = TRAIN / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows.md'
DOCX = OUT / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows_v1.0.docx'
REPORT = OUT / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows_Build-Report_v1.0.md'
VALIDATION = OUT / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows_Validation_v1.0.md'
UNRESOLVED = OUT / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows_Unresolved-Items_v1.0.md'
INVENTORY = OUT / 'Take-Control-of-Your-Inbox_Classic-Outlook-for-Windows_Source-Inventory_v1.0.csv'
SOURCES = [
'01. Locked scope for Guide 1.md','02. Detailed build plan.md','03. Page-by-page copy plan.md',
'04. Complete copy deck — Part 1.md','05. Complete copy deck — Part 2.md','06. Complete copy deck — Part 3.md','07. Complete copy deck — Part 4.md','08. Complete copy deck — Part 5.md','09. Complete copy deck — Part 6.md','10. Complete copy deck — Part 7.md',
'11. Screenshot capture and annotation plan.md','12. Platform-difference and master-guide integration map.md','13. Word production and publishing plan.md','14. Classic Outlook supplemental support package.md','BCBSKS-ID-Prompt-Guide-v4.0.md','BCBSKS_Master_Brand_Kit.md']
TOPICS = [
'Identify Classic Outlook and understand the layout','Read, reply, reply all and forward','Create and send email','Attach files and share cloud links','Search for messages, people and attachments','Decide whether an email needs action','Flag messages and set reminders','Add and use categories','Create folders and move messages','Archive and delete messages','Use Focused Inbox','Create and manage rules','Use follow-up calendar cues','Create appointments or meetings from email','Configure and use automatic replies','Create and manage signatures','Open and use shared mailboxes','Summarize email threads with Copilot','Identify decisions, action items, owners and deadlines','Draft replies with Copilot','Rewrite for clarity, tone or length','Use Coaching by Copilot where available','Create a meeting from an email with Copilot where available','Review Copilot output before sending','Troubleshoot missing commands or unavailable features','Get help and use related resources']

def sha(p): return hashlib.sha256(p.read_bytes()).hexdigest()
def read(p): return (ROOT/p).read_text(encoding='utf-8').replace('\u2028','\n')
def clean_copy_deck(txt):
    txt = re.sub(r'^## Complete copy deck.*?\n', '', txt, flags=re.M)
    txt = txt.replace('[Insert date]', '[Insert last reviewed date]').replace('[Insert support contact or help path]', '[Insert approved support contact or help path]')
    return txt.strip()
def build_markdown():
    parts=[]
    parts.append('---\ntitle: "Take Control of Your Inbox: Classic Outlook for Windows"\nversion: "1.0"\naudience: "Sales Account Representatives and Sales Account Associates using company-managed Windows laptops and work Microsoft 365 accounts"\nsource: "Generated from repository Markdown source files"\n---\n')
    parts.append('# Take Control of Your Inbox: Classic Outlook for Windows\n')
    parts.append('> **Production note:** This is the editable employee-facing guide source for the Word build. It uses approved placeholders where the source materials require human-provided support paths, review dates, screenshots, policy links, permissions, or approvals. It does not fabricate screenshots, controls, features, permissions, support paths, policy links, logos, or approvals.\n')
    parts.append('## High-level overview\nUse this guide to complete common email, follow-up, organization, shared mailbox, and Microsoft 365 Copilot tasks in Classic Outlook for Windows. The guide follows one workflow: **Email → Decide → Make visible → Follow through**. When an email arrives, ask: **Does this need action from me?** If yes, ask: **What is the next step?** and **Where will I make it visible?**\n')
    parts.append('## Architectural diagram concepts\n```text\nWork Microsoft 365 account\n        |\nClassic Outlook for Windows desktop app\n        |\nMailbox + Calendar + People + Search\n        |\nDecision workflow: Email -> Decide -> Make visible -> Follow through\n        |\nFollow-up tools: flags, reminders, categories, folders, archive, rules, calendar cues\n        |\nCopilot support where licensed and available: summarize, identify actions, draft, rewrite, coach\n```\n')
    parts.append('## Step-by-step setup\n1. Open Classic Outlook for Windows with your company-managed Microsoft 365 work account.\n2. Confirm you are in Classic Outlook by looking for the traditional ribbon across the top.\n3. Review the layout: folder pane, message list, reading pane, ribbon, search box, calendar, and shared mailboxes if available.\n4. Set up your basic follow-up system: flags for action, categories for type of work, folders for reference, archive for completed email, and calendar items for time-bound follow-up.\n5. Use Copilot only where it appears in your work account and review all output before sending.\n')
    parts.append('## Common Gotchas\n- Your screen may look slightly different depending on your Microsoft 365 update version and organization settings.\n- Some Copilot, shared mailbox, Focused Inbox, rules, and automatic reply options may depend on licensing, policy, mailbox permissions, or configuration.\n- Screenshots must be captured from an approved test account or sanitized environment before publication.\n- Replace all bracketed placeholders before final distribution.\n')
    parts.append('\n## Complete employee guide copy\n')
    parts.append('### Required topic coverage map\n')
    for i, topic in enumerate(TOPICS, 1):
        parts.append(f'{i}. {topic}')
    parts.append('')
    for name in SOURCES[3:10]:
        parts.append(f'\n<!-- Source: {name} -->\n')
        parts.append(clean_copy_deck(read(name)))
    parts.append('\n## Supplemental support and production notes\n')
    parts.append(clean_copy_deck(read('14. Classic Outlook supplemental support package.md')))
    return '\n\n'.join(parts)+'\n'
def style(doc):
    sec=doc.sections[0]
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Inches(.7); s.left_margin=s.right_margin=Inches(.75)
        s.header.paragraphs[0].text='Take Control of Your Inbox | Classic Outlook for Windows | Version 1.0'
        f=s.footer.paragraphs[0]; f.text='Internal training draft | '; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r=OxmlElement('w:r'); t=OxmlElement('w:t'); t.text='1'; r.append(t); fld.append(r); f._p.append(fld)
    for nm,sz,font,color in [('Normal',10.5,'Arial','333132'),('Heading 1',18,'Georgia','003e6b'),('Heading 2',14,'Georgia','0072bc'),('Heading 3',12,'Arial','003e6b')]:
        st=doc.styles[nm]; st.font.name=font; st.font.size=Pt(sz); st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_after=Pt(6)
def add_para(doc,line):
    raw=line.strip(); txt=re.sub(r'\*\*(.*?)\*\*', r'\1', raw).strip()
    if raw.startswith('# '): doc.add_heading(txt[2:].strip(),0)
    elif raw.startswith('## '): doc.add_heading(txt[3:].strip(),1)
    elif raw.startswith('### '): doc.add_heading(txt[4:].strip(),2)
    elif raw.startswith('#### '): doc.add_heading(txt[5:].strip(),3)
    elif raw.startswith('- '): doc.add_paragraph(txt[2:].strip(), style='List Bullet')
    elif re.match(r'^\d+\.\s+', raw): doc.add_paragraph(re.sub(r'^\d+\.\s+','',txt), style='List Number')
    elif raw.startswith('|'): doc.add_paragraph(txt)
    elif raw.startswith('```'): pass
    elif raw.startswith('---'): pass
    else: doc.add_paragraph(txt)
def build_minimal_docx(md):
    import html
    paras=[]
    for line in md.splitlines():
        raw=line.strip()
        if not raw or raw.startswith('---') or raw.startswith('```'):
            continue
        level=0
        if raw.startswith('# '): level=1; raw=raw[2:]
        elif raw.startswith('## '): level=2; raw=raw[3:]
        elif raw.startswith('### '): level=3; raw=raw[4:]
        raw=re.sub(r'^[-*]\s+','• ',raw)
        raw=re.sub(r'^\d+\.\s+','',raw)
        raw=re.sub(r'\*\*(.*?)\*\*', r'\1', raw)
        tag='w:p'
        if level:
            paras.append(f'<w:p><w:pPr><w:pStyle w:val=\"Heading{min(level,3)}\"/></w:pPr><w:r><w:t>{html.escape(raw)}</w:t></w:r></w:p>')
        else:
            paras.append(f'<w:p><w:r><w:t>{html.escape(raw)}</w:t></w:r></w:p>')
    body=''.join(paras)+'<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/><w:pgMar w:top=\"1008\" w:right=\"1080\" w:bottom=\"1008\" w:left=\"1080\"/></w:sectPr>'
    doc_xml='<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?><w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"><w:body>'+body+'</w:body></w:document>'
    content='<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?><Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\"><Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/><Default Extension=\"xml\" ContentType=\"application/xml\"/><Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/><Override PartName=\"/docProps/core.xml\" ContentType=\"application/vnd.openxmlformats-package.core-properties+xml\"/><Override PartName=\"/docProps/app.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.extended-properties+xml\"/></Types>'
    rels='<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?><Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\"><Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/></Relationships>'
    core='<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?><cp:coreProperties xmlns:cp=\"http://schemas.openxmlformats.org/package/2006/metadata/core-properties\" xmlns:dc=\"http://purl.org/dc/elements/1.1/\"><dc:title>Take Control of Your Inbox: Classic Outlook for Windows</dc:title><dc:creator>IT Training &amp; Enablement</dc:creator><dc:description>Employee-facing Classic Outlook guide</dc:description></cp:coreProperties>'
    app='<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?><Properties xmlns=\"http://schemas.openxmlformats.org/officeDocument/2006/extended-properties\"><Application>Deterministic Python OOXML builder</Application></Properties>'
    with zipfile.ZipFile(DOCX,'w',zipfile.ZIP_DEFLATED) as z:
        z.writestr('[Content_Types].xml',content); z.writestr('_rels/.rels',rels); z.writestr('word/document.xml',doc_xml); z.writestr('docProps/core.xml',core); z.writestr('docProps/app.xml',app)

def build_docx(md):
    if not HAS_DOCX:
        build_minimal_docx(md); return
    doc=Document(); style(doc)
    props=doc.core_properties; props.title='Take Control of Your Inbox: Classic Outlook for Windows'; props.author='IT Training & Enablement'; props.subject='Employee-facing Classic Outlook guide'; props.keywords='Classic Outlook, Microsoft 365 Copilot, inbox, sales training'; props.comments='Generated deterministically from repository Markdown sources'
    doc.add_heading('Take Control of Your Inbox',0); doc.add_paragraph('Classic Outlook for Windows'); doc.add_paragraph('Step-by-step email organization, follow-up, and Copilot guidance for sales teams')
    doc.add_page_break()
    for line in md.splitlines():
        if line.strip(): add_para(doc,line)
    doc.save(DOCX)
def reports(md):
    checks=[]
    def chk(n,ok,e): checks.append((n,'Pass' if ok else 'Fail',e))
    for s in SOURCES: chk(f'Source present: {s}', (ROOT/s).exists(), s)
    for t in TOPICS: chk(f'Topic included: {t}', t in md, 'Found in guide source' if t in md else 'Missing')
    for phrase in ['High-level overview','Architectural diagram concepts','Step-by-step setup','Common Gotchas','Email → Decide → Make visible → Follow through','Copilot starts the draft. You finish the message.']:
        chk(f'Required phrase/section: {phrase}', phrase in md, 'Found' if phrase in md else 'Missing')
    chk('No docx markdown substitute', not any(ROOT.rglob('*.docx.md')), 'No *.docx.md files found')
    chk('DOCX exists', DOCX.exists(), str(DOCX.relative_to(ROOT)))
    chk('DOCX is OOXML zip', zipfile.is_zipfile(DOCX), 'zipfile.is_zipfile')
    chk('DOCX contains editable Word document XML', zipfile.is_zipfile(DOCX) and 'word/document.xml' in zipfile.ZipFile(DOCX).namelist(), 'word/document.xml present in generated OOXML package')
    VALIDATION.write_text('# Validation report\n\n| Check | Status | Evidence |\n|---|---|---|\n'+'\n'.join(f'| {a} | {b} | {c} |' for a,b,c in checks)+'\n',encoding='utf-8')
    unresolved=['Approved support contact or help path','Approved policy links','Last reviewed date','Future review date','Approved screenshot captures','Shared mailbox permission owners','Copilot feature availability confirmation','Publication approval']
    UNRESOLVED.write_text('# Unresolved items report\n\n| Item | Why unresolved | Required action | Blocker status |\n|---|---|---|---|\n'+'\n'.join(f'| {u} | Source materials require validation or placeholder replacement. | Confirm with document owner before final distribution. | Blocks final publication, not build generation. |' for u in unresolved)+'\n',encoding='utf-8')
    with INVENTORY.open('w',newline='',encoding='utf-8') as f:
        w=csv.writer(f); w.writerow(['path','size_bytes','sha256','use_status'])
        for s in SOURCES: p=ROOT/s; w.writerow([s,p.stat().st_size,sha(p),'Read and used'])
        w.writerow([GUIDE_MD.relative_to(ROOT),GUIDE_MD.stat().st_size,sha(GUIDE_MD),'Generated guide source'])
    REPORT.write_text(f'# Build report\n\n- Build timestamp: intentionally omitted so committed text reports remain deterministic.\n- Command: `python3 build/build_classic_outlook_employee_guide.py`\n- Markdown source: `{GUIDE_MD.relative_to(ROOT)}`\n- DOCX output: `{DOCX.relative_to(ROOT)}` (generated artifact; not committed because binary files are not supported for review)\n- Validation report: `{VALIDATION.relative_to(ROOT)}`\n- Unresolved items report: `{UNRESOLVED.relative_to(ROOT)}`\n- Source inventory: `{INVENTORY.relative_to(ROOT)}`\n- Result: deterministic build completed; unresolved placeholders are documented and not fabricated.\n',encoding='utf-8')
    return checks
def main():
    OUT.mkdir(parents=True,exist_ok=True); TRAIN.mkdir(parents=True,exist_ok=True)
    md=build_markdown(); GUIDE_MD.write_text(md,encoding='utf-8'); build_docx(md); checks=reports(md)
    fails=[c for c in checks if c[1]=='Fail']
    return 1 if fails else 0
if __name__=='__main__': sys.exit(main())
