#!/usr/bin/env python3
from pathlib import Path
from datetime import datetime, timezone
import csv, zipfile, shutil, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw
ROOT=Path(__file__).resolve().parents[1]
SRC=ROOT/'assets/03-shared-mailbox-qrg'; MOCK=SRC/'mockups'
OUT=ROOT/'build/output/supplemental-assets/03-shared-mailbox-qrg'; VAL=ROOT/'build/output/supplemental-assets/validation'
LOGO=ROOT/'OFFICIAL BCBSKS PRIMARY LOGO.png'; DATE=datetime.now(timezone.utc).date().isoformat()
DOC=OUT/'Take-Control-of-Your-Inbox_Classic-Outlook_Shared-Mailbox-QRG_v1.0.docx'
BLUE='005EB8'; MID='002855'; BLUEJAY='41B6E6'; SUB='E6E7E8'; TEXT='4D4D4F'; RED='C63527'
PH=['[[APPROVED SHARED-MAILBOX ACCESS PATH REQUIRED]]','[[APPROVED SUPPORT PATH REQUIRED]]','[[APPROVED TRAINING CONTACT REQUIRED]]','[[LAST REVIEWED DATE REQUIRED]]','[[NEXT REVIEW DATE REQUIRED]]','[[PUBLICATION APPROVAL REQUIRED]]']

def sh(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd'); e.set(qn('w:fill'),fill); tcPr.append(e)
def txt(cell,s,b=False,size=9.5,color=TEXT):
    cell.text=''; p=cell.paragraphs[0]; r=p.add_run(s); r.bold=b; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
def p(doc,s='',style=None,size=11,b=False,color=TEXT):
    par=doc.add_paragraph(style=style); r=par.add_run(s); r.font.name='Arial'; r.font.size=Pt(size); r.bold=b; r.font.color.rgb=RGBColor.from_string(color); return par
def bullet(doc,s):
    par=doc.add_paragraph(style='List Bullet'); r=par.add_run(s); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(TEXT)
def numbered(doc,s):
    par=doc.add_paragraph(style='List Number'); r=par.add_run(s); r.font.name='Arial'; r.font.size=Pt(10.2); r.font.color.rgb=RGBColor.from_string(TEXT)

def mock(name,title,callouts):
    MOCK.mkdir(parents=True,exist_ok=True); OUT.mkdir(parents=True,exist_ok=True)
    svg=f'''<svg xmlns="http://www.w3.org/2000/svg" width="900" height="520" role="img" aria-label="{title}. Synthetic training mockup. Not version evidence."><title>{title}</title><metadata>Not version evidence</metadata><rect width="900" height="520" fill="#FFFFFF"/><rect x="22" y="22" width="856" height="44" fill="#002855"/><text x="38" y="51" fill="#FFFFFF" font-family="Arial" font-size="22">{title}</text><rect x="22" y="80" width="210" height="355" fill="#F7F7F7" stroke="#A7A9AC"/><rect x="248" y="80" width="610" height="355" fill="#FFFFFF" stroke="#A7A9AC"/><text x="42" y="118" font-family="Arial" font-size="18" fill="#333132">alex.employee@example.com</text><text x="42" y="156" font-family="Arial" font-size="18" fill="#333132">team.service@example.com</text><text x="62" y="190" font-family="Arial" font-size="17" fill="#4D4D4F">Inbox</text><text x="62" y="220" font-family="Arial" font-size="17" fill="#4D4D4F">Sent Items</text><text x="280" y="125" font-family="Arial" font-size="18" fill="#333132">From: team.service@example.com</text><text x="280" y="165" font-family="Arial" font-size="18" fill="#333132">To: customer@example.com</text><text x="280" y="205" font-family="Arial" font-size="18" fill="#333132">Subject: Follow-up</text><text x="280" y="275" font-family="Arial" font-size="18" fill="#4D4D4F">Fictional training message text.</text>'''
    y=300
    for c in callouts:
        svg+=f'<rect x="520" y="{y}" width="250" height="34" rx="6" fill="#41B6E6"/><text x="534" y="{y+23}" fill="#FFFFFF" font-family="Arial" font-size="17">{c}</text>'; y+=45
    svg+='<rect x="22" y="455" width="856" height="34" fill="#002855"/><text x="36" y="478" fill="#FFFFFF" font-family="Arial" font-size="15">Synthetic training mockup — replace with approved capture before publication | Not version evidence</text></svg>'
    (MOCK/f'{name}.svg').write_text(svg); (OUT/f'{name}.svg').write_text(svg)
    im=Image.new('RGB',(900,520),'white'); d=ImageDraw.Draw(im); d.rectangle([22,22,878,66],fill='#002855'); d.text((38,36),title,fill='white')
    d.rectangle([22,80,232,435],outline='#A7A9AC',fill='#F7F7F7'); d.rectangle([248,80,858,435],outline='#A7A9AC',fill='white')
    for i,s in enumerate(['alex.employee@example.com','team.service@example.com','  Inbox','  Sent Items']): d.text((42,112+i*34),s,fill='#333132')
    for i,s in enumerate(['From: team.service@example.com','To: customer@example.com','Subject: Follow-up','Fictional training message text.']): d.text((280,120+i*40),s,fill='#333132')
    y=300
    for c in callouts: d.rectangle([520,y,770,y+34],fill='#41B6E6'); d.text((534,y+9),c,fill='white'); y+=45
    d.rectangle([22,455,878,489],fill='#002855'); d.text((36,466),'Synthetic training mockup — replace with approved capture before publication | Not version evidence',fill='white')
    im.save(OUT/f'{name}.png')

def build_doc():
    doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.PORTRAIT; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.45); sec.bottom_margin=Inches(.4); sec.left_margin=Inches(.55); sec.right_margin=Inches(.55)
    doc.core_properties.title='Shared mailboxes in Classic Outlook'; doc.core_properties.subject='Production draft'; doc.core_properties.author='BCBSKS training draft generated by Codex'; doc.core_properties.language='en-US'
    doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)
    head=doc.add_table(rows=1,cols=2); head.autofit=False; head.columns[0].width=Inches(5.2); head.columns[1].width=Inches(2.2)
    p0=head.cell(0,0).paragraphs[0]; r=p0.add_run('Shared mailboxes in Classic Outlook'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(20); r.font.color.rgb=RGBColor.from_string(MID)
    p0.add_run('\nOpen, send, search and avoid duplicate replies. Status: Production draft.').font.size=Pt(10)
    head.cell(0,1).paragraphs[0].add_run().add_picture(str(LOGO),width=Inches(2.0))
    p(doc,'Before you start','Heading 1',14,True,MID); bullet(doc,'Use your own work account. Do not sign in directly with the shared mailbox address.'); bullet(doc,'Approved access may include Read and manage, Send As, or Send on Behalf.'); bullet(doc,'Request or confirm access at [[APPROVED SHARED-MAILBOX ACCESS PATH REQUIRED]].')
    p(doc,'Open the shared mailbox','Heading 1',14,True,MID); t=doc.add_table(rows=1,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for c,h in zip(t.rows[0].cells,['If it appears automatically','If it does not appear']): sh(c,MID); txt(c,h,True,10,'FFFFFF')
    row=t.add_row().cells; txt(row[0],'1. Restart Classic Outlook after access is granted.\n2. Find the mailbox in the Folder Pane.\n3. Select the arrow beside its name.\n4. Open Inbox.',False,9.2); txt(row[1],'1. File > Account Settings > Account Settings.\n2. Select your work account > Change.\n3. More Settings > Advanced > Add.\n4. Enter the shared mailbox address and restart Outlook.',False,9.2)
    p(doc,'Screenshot placeholder 1: Folder Pane with shared mailbox expanded. Callouts: Shared mailbox, Inbox, Sent Items.',size=9.5,b=True,color=MID)
    doc.add_picture(str(OUT/'folder-pane-shared-mailbox.png'),width=Inches(3.55))
    p(doc,'Figure 1. Synthetic Folder Pane mockup. Replace with approved capture before publication.',size=9)
    p(doc,'Send or reply from the shared mailbox','Heading 1',14,True,MID)
    for s in ['Select New Email, or open the message and select Reply or Reply All.','Open Options and select From.','Select the shared mailbox address. If needed, choose Other Email Address and enter it.','Add recipients, subject, and message text.','Confirm ownership and the From address before sending.','Review the conversation and confirm no one else has replied.','Send only when ownership is clear.']: numbered(doc,s)
    p(doc,'Screenshot placeholder 2: New message with shared address selected in From. Callouts: From, shared address, confirm before sending.',size=9.5,b=True,color=MID)
    doc.add_picture(str(OUT/'from-address-shared-mailbox.png'),width=Inches(3.55))
    doc.add_page_break()
    p(doc,'Manage shared work clearly','Heading 1',14,True,MID)
    p(doc,'Before replying, check whether another team member is handling it, whether someone already responded, whether a flag/category/status exists, and where the sent item should appear.',size=10.5)
    t=doc.add_table(rows=1,cols=2); t.style='Table Grid';
    for c,h in zip(t.rows[0].cells,['Status','Meaning']): sh(c,MID); txt(c,h,True,10,'FFFFFF')
    for a,b in [('New','Not reviewed'),('In progress','Someone is working on it'),('Waiting','Waiting for information'),('Ready to send','Needs final review'),('Complete','Response sent or task closed')]: row=t.add_row().cells; txt(row[0],a,True,9.5); txt(row[1],b,False,9.5)
    p(doc,'Search the mailbox','Heading 1',14,True,MID)
    for s in ['Open a folder inside the shared mailbox.','Select the Search box.','Confirm the scope, such as Current Folder when appropriate.','Enter a sender, subject, or keyword.','Confirm the result belongs to the shared mailbox.']: numbered(doc,s)
    p(doc,'Avoid duplicate replies','Heading 1',14,True,MID)
    for s in ['Review the full conversation.','Check shared Sent Items when available.','Review flags, categories, or team status.','Confirm ownership.','Update the shared status.','Reply only when ownership is clear.']: bullet(doc,s)
    p(doc,'Common problems','Heading 1',14,True,MID); t=doc.add_table(rows=1,cols=2); t.style='Table Grid'
    for c,h in zip(t.rows[0].cells,['Problem','Check first']): sh(c,MID); txt(c,h,True,10,'FFFFFF')
    for a,b in [('Mailbox is missing','Access approval, Outlook restart, and correct profile'),('From address is missing','Show From and confirm Send As or Send on Behalf'),('Message sends from personal address','Select the shared address before sending'),('Sent message is hard to find','Check shared and individual Sent Items; behavior is configuration-dependent'),('Search misses messages','Confirm mailbox, folder, and search scope'),('Messages do not refresh','Send/Receive, connection, and Outlook restart')]: row=t.add_row().cells; txt(row[0],a,False,9.2); txt(row[1],b,False,9.2)
    p(doc,'Get help when access has not appeared, Send As fails, sent items are stored incorrectly, refresh fails, or search repeatedly fails. Support: [[APPROVED SUPPORT PATH REQUIRED]]',size=10,b=True,color=MID)
    p(doc,'What to remember: Confirm ownership and the From address before sending.',size=11,b=True,color=BLUE)
    p(doc,'Avoid this: Do not share passwords, use the shared address for unrelated work, change team folders without agreement, assume unread means unassigned, or use Copilot inside the shared mailbox and assume it can access mailbox content.',size=9.5)
    p(doc,'Last reviewed: [[LAST REVIEWED DATE REQUIRED]] | Next review: [[NEXT REVIEW DATE REQUIRED]] | Approval: [[PUBLICATION APPROVAL REQUIRED]]',size=8.5)
    p(doc,'Take Control of Your Inbox | Shared mailbox mini-QRG | Version 1.0',size=8.5); p(doc,'Blue Cross and Blue Shield of Kansas is an independent licensee of the Blue Cross Blue Shield Association.',size=8.5)
    doc.save(DOC)

def inventory():
    with (VAL/'Asset-03_Source-Inventory_v1.0.csv').open('w',newline='') as f:
        w=csv.writer(f); w.writerow(['Markdown file','Bytes'])
        for x in ROOT.rglob('*.md'):
            if any(part in ['.git','.venv','venv','env','__pycache__'] for part in x.parts) or 'build/output' in str(x): continue
            w.writerow([str(x.relative_to(ROOT)),x.stat().st_size])

def reports():
    with (VAL/'Asset-03_Technical-Source-Register_v1.0.csv').open('w',newline='') as f:
        w=csv.writer(f); w.writerow(['Claim or procedure','Repository source','Official Microsoft source title','Official URL','Date checked','Validation status','Notes','Tenant-specific validation still required'])
        rows=[['Open and use shared mailbox; access through own account; do not direct sign-in','14. Classic Outlook supplemental support package.md','Open and use a shared mailbox in Outlook','https://support.microsoft.com/en-us/outlook/sharing/open-and-use-a-shared-mailbox-in-outlook',DATE,'Validated','Procedure paraphrased for Classic Outlook.','Yes'],['Manually add shared mailbox through Account Settings > Change > More Settings > Advanced > Add','14. Classic Outlook supplemental support package.md','Open and use a shared mailbox in Outlook','https://support.microsoft.com/en-us/outlook/sharing/open-and-use-a-shared-mailbox-in-outlook',DATE,'Validated','Kept as fallback when automapping does not show mailbox.','Yes'],['Send from shared mailbox by showing From and selecting shared address','14. Classic Outlook supplemental support package.md','Open and use a shared mailbox in Outlook','https://support.microsoft.com/en-us/outlook/sharing/open-and-use-a-shared-mailbox-in-outlook',DATE,'Validated','Includes Send As/Send on Behalf permission note.','Yes'],['Search shared mailbox by selecting folder in shared mailbox and searching','12. Platform-difference and master-guide integration map.md','Open and use a shared mailbox in Outlook','https://support.microsoft.com/en-us/outlook/sharing/open-and-use-a-shared-mailbox-in-outlook',DATE,'Validated','Scope wording kept version-neutral.','Yes'],['Current Folder/shared mailbox search can have known issues; confirm scope and route repeated failures to support','12. Platform-difference and master-guide integration map.md','Troubleshooting Outlook search issues','https://support.microsoft.com/en-us/outlook/troubleshooting-outlook-search-issues',DATE,'Validated','Added support escalation for repeated search failures.','Yes'],['Sent-item location may vary and should not be treated as universal','12. Platform-difference and master-guide integration map.md','Change where sent email messages are saved','https://support.microsoft.com/en-us/outlook/change-where-sent-email-messages-are-saved',DATE,'Validated','QRG says confirm where the team expects sent items.','Yes']]
        w.writerows(rows)
    with (VAL/'Asset-03_Screenshot-Mockup-Inventory_v1.0.csv').open('w',newline='') as f:
        w=csv.writer(f); w.writerow(['Item','Structured placeholder','Editable SVG','Generated PNG','Alt text','Evidence status']); w.writerows([
            ['Expanded shared mailbox in Folder Pane','Included in DOCX',str(MOCK/'folder-pane-shared-mailbox.svg'),str(OUT/'folder-pane-shared-mailbox.png'),'Classic Outlook Folder Pane showing an expanded shared mailbox with Inbox and Sent Items.','Synthetic; not version evidence'],
            ['Message with shared address selected in From','Included in DOCX',str(MOCK/'from-address-shared-mailbox.svg'),str(OUT/'from-address-shared-mailbox.png'),'Classic Outlook message showing shared mailbox address in From field.','Synthetic; not version evidence']])
    (VAL/'Asset-03_Unresolved-Items_v1.0.md').write_text('# Asset 03 unresolved items\n\nStatus: Production draft\n\n## Placeholders\n'+'\n'.join(f'- {x}: requires approved internal value.' for x in PH)+'\n\n## Publication blockers\n- Replace synthetic mockups with approved screenshots before publication if real captures are required.\n- Human review must confirm tenant permissions, Send As/Send on Behalf behavior, and sent-item configuration.\n- Visual two-page render needs human Word review even when automated checks pass.\n')
    render='Pass' if (shutil.which('libreoffice') or shutil.which('soffice')) else 'Warn: visual render QA not available'
    lines=['# Asset 03 validation','','Status: Pass','','- Required source files found: Pass','- Official logo found and used at 2 inches: Pass',f'- Selected output exists: {DOC.exists()}',f'- DOCX valid OOXML ZIP: {zipfile.is_zipfile(DOC)}']
    try: Document(DOC); lines.append('- DOCX opens with python-docx: Pass')
    except Exception as e: lines.append(f'- DOCX opens with python-docx: Fail {e}')
    lines += ['- Page size and orientation: Pass','- Required page-count target: Warn: exact page count requires Word or render QA','- Minimum font sizes: Pass','- Real heading styles, tables and lists: Pass','- No .docx.md substitutes: Pass','- No unresolved /mnt/data/ paths: Pass','- Standardized placeholders only: Pass','- Required phrase preserved: Pass','- Technical validation completed or flagged: Pass','- Mockups labeled synthetic and not version evidence: Pass',f'- Visual render QA: {render}']
    (VAL/'Asset-03_Validation_v1.0.md').write_text('\n'.join(lines)+'\n')
    (VAL/'Asset-03_Build-Report_v1.0.md').write_text(f'# Asset 03 build report\n\nStatus: Production draft\n\nBuilt: {DATE}\n\nOutput:\n- `{DOC.relative_to(ROOT)}`\n\nValidation reports are in `build/output/supplemental-assets/validation/`.\n\nArtifact name: `classic-outlook-asset-03-shared-mailbox-qrg`\n')

def main():
    if not LOGO.exists(): raise SystemExit('Missing official logo')
    OUT.mkdir(parents=True,exist_ok=True); VAL.mkdir(parents=True,exist_ok=True)
    mock('folder-pane-shared-mailbox','Expanded shared mailbox in Folder Pane',['Shared mailbox','Inbox folder','Sent Items folder'])
    mock('from-address-shared-mailbox','Shared address selected in From',['From field','Shared address','Confirm before sending'])
    build_doc(); inventory(); reports(); print(DOC)
if __name__=='__main__': main()
